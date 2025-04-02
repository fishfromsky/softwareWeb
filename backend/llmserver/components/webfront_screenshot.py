import subprocess
import os
import time
from selenium.webdriver.support import expected_conditions as EC
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from tenacity import retry, stop_after_attempt, wait_exponential
from openai import OpenAI
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
import json
import sys
import shutil
import re
from utils import add_multi_level, add_manual, add_pager_header
import threading
from natsort import natsorted
import base64
from PIL import Image, ImageDraw, ImageFont
from graphviz import Digraph
from concurrent.futures import ThreadPoolExecutor, as_completed
import queue
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.append(BASE_DIR)

from backend import settings


MEDIUM_PATH = os.path.join(BASE_DIR, "medium")

# Vue 项目路径
vue_project_path = os.path.join(os.path.dirname(BASE_DIR), "webfront")

chrome_user_path = os.path.join(BASE_DIR, "chrome_user_data")


client = OpenAI(
    api_key=settings.LLM_API_KEY,
    base_url=settings.LLM_BASE_URL
)


def exponential_search(json_data):
    # 随机选择一个起始点
    keys = json_data.keys()
    for key in keys:
        next_key = key
        if json_data[key] == False:
            return next_key
    return -1


def write_env_file(path, username, datetime, colors):
    env_file_path = os.path.join(path, "user.json")

    if isinstance(colors, str):
        colors_list = colors.split(',')
    else:
        colors_list = colors

    user_dict = {
        "username": username,
        "datetime": datetime,
        "colors": colors_list
    }

    with open(env_file_path, "w", encoding="utf-8") as json_file:
        json.dump(user_dict, json_file, ensure_ascii=False, indent=2)


def modify_port_file_start(path):  # 将对应端口改为True
    with open(path, "r", encoding="utf-8") as fr:
        json_data = json.load(fr)
        fr.close()

    port = exponential_search(json_data)
    json_data[port] = True

    with open(path, "w", encoding="utf-8") as fw:
        json.dump(json_data, fw)
        fw.close()

    return port


def modify_port_file_end(path, port):
    with open(path, "r", encoding="utf-8") as fr:
        json_data = json.load(fr)
        fr.close()

    json_data[port] = False

    with open(path, "w", encoding="utf-8") as fw:
        json.dump(json_data, fw)
        fw.close()


def copy_file(file_path1, file_path2):
    shutil.copytree(file_path1, file_path2)


def kill_process_by_port(port):  # 强制终止进程 (限于Linux系统)
    try:
        # 查找占用指定端口的进程ID
        command = f"lsof -i :{port} -t"
        pid = subprocess.check_output(command, shell=True)
        pid = int(pid.strip())

        # 尝试终止进程
        os.kill(pid, 15)  # 先尝试发送SIGTERM信号请求正常终止
        print(f"已向PID {pid}发送终止信号，等待进程正常结束...")

        # 确认进程是否已经结束，如果没有，则强制终止
        try:
            os.kill(pid, 0)  # 这个操作会抛出异常如果进程不存在
            os.kill(pid, 9)  # 如果进程还在运行，发送SIGKILL信号强制结束
            print(f"已强制终止PID {pid}")
        except OSError:
            print(f"PID {pid} 已正常结束")

    except subprocess.CalledProcessError:
        print(f"找不到端口号为 {port} 的进程")
    except ValueError:
        print("解析PID时出现错误")
    except OSError as e:
        print(f"无法终止进程: {e}")


# 启动 Vue 项目
def start_vue_project(username, datetime, colors):

    virtual_vue_path = os.path.join(BASE_DIR, "virtual_vue_project", username, datetime, "webfront")
    copy_file(vue_project_path, virtual_vue_path)

    # 进入 Vue 项目目录并启动开发服务器
    write_env_file(virtual_vue_path, username, datetime, colors)
    json_path = os.path.join(BASE_DIR, "llmserver", "port.json")

    port = modify_port_file_start(json_path)

    print("正在编译")
    process = subprocess.Popen(
        ["npm", "install"],
        cwd=virtual_vue_path,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True
    )
    while True:
        output = process.stdout.readline()
        if output == "" and process.poll() is not None:
            break
        if output:
            print(output.strip())

    print("编译完成")

    process = subprocess.Popen(
        ["npm", "run", "serve", "--", "--port", str(port)],
        cwd=virtual_vue_path,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True
    )
    success = False
    while True:
        output = process.stdout.readline()
        if output == "" and process.poll() is not None:
            break
        if output:
            print(output.strip())
            if "App running at:" in output or "Local:   http://localhost:" in output:
                success = True
                print("Vue 项目已成功启动！")
                break

    if not success:
        print("Vue 项目启动失败！")

    return process, port, virtual_vue_path

# 关闭 Vue 项目
def stop_vue_project(process, port, virtual_vue_path):
    process.terminate()
    process.kill()

    json_path = os.path.join(BASE_DIR, "llmserver", "port.json")

    modify_port_file_end(json_path, port)

    kill_process_by_port(port)

    virtual_vue_path = os.path.dirname(virtual_vue_path)
    shutil.rmtree(virtual_vue_path)
    print("Vue 项目已关闭。")

def generate_config(PLATFORM, username, datetime):
    try:
        with open(os.path.join(BASE_DIR, "medium", username, datetime, "menu.json"), "r", encoding="utf-8") as f:
            menu = json.load(f)
        json_str = json.dumps(menu, ensure_ascii=False)
        MESSAGE = [{"role": "system", "content": "You are a helpful programmer and product manager"}]
        question = f"""
        现在有一个名为{PLATFORM}的后台管理系统，其侧边栏为{json_str}，请为其设计如下的配置信息：

        1. "系统简介"：请撰写系统的整体介绍，重点包括系统背景、开发目的与设计理念。需结合实际应用场景，阐述系统诞生的原因、所要解决的问题、目标用户及其使用方式。设计理念部分应聚焦于系统架构构思、用户体验、性能、安全性或可维护性等方面的思考与权衡。
        2. "功能描述"：对侧边栏每个父目录及其子目录进行一句话功能描述。
        3. "功能描述"内容以嵌套Object形式返回，第一层的每个Key为父目录名称，Value是一个子Object，该子Object中每个Key为子目录名称，value为对应的功能描述。
        4. "功能描述"的Object中的每个子Object的每个键值对之间必须以英文逗号“,”分隔，禁止使用分号。
        5. 所有Object的Key必须使用双引号引起来。
        6. 生成的所有配置信息为纯文字说明，不需要换行。
        7. 返回内容为一个JSON格式的Object，以“{{”开头，“}}”结尾，中间不要插入任何解释或说明。

        ### 请从这里开始生成配置内容 ###
        """
        query = {"role": "user", "content": question}
        MESSAGE.append(query)
        completion = client.chat.completions.create(
            model="qwen-coder-plus-latest",
            messages=MESSAGE
        )
        data_dict = json.loads(completion.model_dump_json())
        content = data_dict["choices"][0]["message"]["content"]
        content = json.loads(content, strict=False)
        return content
    except Exception as e:
        raise e



def draw_annotations(image_path, elements):
    """使用百分比位置绘制标注"""
    try:
        img = Image.open(image_path)
        draw = ImageDraw.Draw(img)
        
        # 获取图片尺寸
        img_width, img_height = img.size
        
        # 直接使用simhei.ttf字体
        font_size = 26
        simhei_path = "/usr/share/fonts/simhei.ttf"
        
        try:
            # 加载simhei.ttf字体
            number_font = ImageFont.truetype(simhei_path, font_size)
            print(f"成功加载字体: {simhei_path}")
        except Exception as e:
            print(f"加载simhei.ttf失败: {str(e)}")
            # 如果加载失败，使用默认字体
            number_font = ImageFont.load_default()
                
        desc_font = number_font  # 使用相同的字体            
        for i, element in enumerate(elements):
            pos = element["position"]
            
            # 将百分比转换回像素坐标
            x = int((pos["x"] / 100) * img_width)
            y = int((pos["y"] / 100) * img_height)
            width = int((pos["width"] / 100) * img_width)
            height = int((pos["height"] / 100) * img_height)
            
            # 绘制矩形框
            draw.rectangle(
                [(x-20, y-10), (x + width+10, y + height+10)],
                outline="red",
                width=5
            )
            
            # 在框的左上角添加序号
            draw.text(
                (x-10, max(0, y - 25)),
                str(i + 1),
                fill="red",
                font=number_font
            )
            
            # 在底部添加说明文字
            note_y = img_height - 30 * (len(elements) - i)
            draw.text(
                (10, note_y),
                f"{i+1}. {element['description'][:50]}...",
                fill="red",
                font=desc_font
            )
        
        annotated_path = image_path.replace(".png", "_annotated.png")
        img.save(annotated_path)
        return annotated_path
        
    except Exception as e:
        print(f"绘制标注时出错: {str(e)}")
        return None
def analyze_page_elements(driver, client):
    """分析页面元素并获取相对位置（百分比），只查找表格和按钮，并获取HTML描述"""
    elements_info = []
    
    # 获取页面可视区域的尺寸
    viewport_width = driver.execute_script("return document.documentElement.clientWidth")
    viewport_height = driver.execute_script("return document.documentElement.clientHeight")
    
    selectors = {
        "表格": ".el-table",
        "表单": "form",
        "功能按钮": ".el-button:not(.el-table *)",
    }
    
    for element_type, selector in selectors.items():
        elements = driver.find_elements(By.CSS_SELECTOR, selector)
        for element in elements:
            try:
                # 获取元素的HTML代码
                html_content = driver.execute_script("return arguments[0].outerHTML;", element)
                
                # 根据元素类型生成不同的提示词
                if element_type == "表格":
                    prompt = f"""
                    分析这个表格的HTML代码，用10-20字描述它的主要功能：
                    {html_content}
                    只需要返回描述文字，不需要任何解释。
                    """
                else:  # 按钮类型
                    prompt = f"""
                    分析这个按钮的HTML代码，用10-20字描述它的功能：
                    {html_content}
                    只需要返回描述文字，不需要任何解释。
                    """
                
                # 调用AI获取描述
                try:
                    completion = client.chat.completions.create(
                        model="qwen-plus",
                        messages=[{"role": "user", "content": prompt}]
                    )
                    ai_description = completion.choices[0].message.content.strip()
                except Exception as e:
                    print(f"AI描述生成失败: {str(e)}")
                    ai_description = "数据表格" if element_type == "表格" else element.text.strip()
                
                # 获取元素位置信息
                rect = driver.execute_script("""
                    var rect = arguments[0].getBoundingClientRect();
                    var scrollLeft = window.pageXOffset || document.documentElement.scrollLeft;
                    var scrollTop = window.pageYOffset || document.documentElement.scrollTop;
                    var viewWidth = document.documentElement.clientWidth;
                    var viewHeight = document.documentElement.clientHeight;
                    
                    return {
                        x: ((rect.left + scrollLeft) / viewWidth) * 100,
                        y: ((rect.top + scrollTop) / viewHeight) * 100,
                        width: (rect.width / viewWidth) * 100,
                        height: (rect.height / viewHeight) * 100
                    };
                """, element)
                
                elements_info.append({
                    "type": element_type,
                    "position": {
                        "x": float(rect['x']),
                        "y": float(rect['y']),
                        "width": float(rect['width']),
                        "height": float(rect['height'])
                    },
                    "description": f"{element_type}: {ai_description}"
                })
            except Exception as e:
                print(f"处理元素时出错: {str(e)}")
    
    return elements_info


def take_screenshot(index, driver, IMAGE_PATH):
    script = f"""
        window.dataLoaded = false;
        window.addEventListener("data-loaded_{index}", () => {{
            window.dataLoaded = true;
        }});
    """
    driver.execute_script(script)
    wait = WebDriverWait(driver, 500)
    wait.until(lambda d: d.execute_script("return window.dataLoaded;"))
    time.sleep(10)

    colors_data = driver.execute_script("return window.myColors || null;")
    username_data = driver.execute_script("return window.userName || null;")

    try:
        # 先保存主页面截图
        screenshot_path = os.path.join(IMAGE_PATH, f"{index}.png")
        driver.save_screenshot(screenshot_path)
        print(f"截图已保存为 {screenshot_path}")
        elements = analyze_page_elements(driver, client)
        
        # 在截图上添加标注
        if elements:
            annotated_path = draw_annotations(screenshot_path, elements)
            print(f"已生成标注后的截图: {annotated_path}")
            
        # 查找所有新增按钮
        add_buttons = driver.find_elements(By.CLASS_NAME, "add")
        print(f"找到 {len(add_buttons)} 个新增按钮")

        # 遍历每个新增按钮
        for btn_index, add_button in enumerate(add_buttons, 1):
            try:
                # 检查按钮是否可见和可点击
                if not add_button.is_displayed() or not add_button.is_enabled():
                    print(f"[{index}] 按钮 {btn_index} 不可见或不可点击，跳过")
                    continue

                # 滚动到按钮位置
                driver.execute_script("arguments[0].scrollIntoView(true);", add_button)
                time.sleep(1)

                # 使用 JavaScript 点击按钮
                driver.execute_script("arguments[0].click();", add_button)
                print(f"[{index}] 点击新增按钮 {btn_index}")
                time.sleep(1)

                try:
                    # 等待对话框出现
                    dialog = WebDriverWait(driver, 5).until(
                        EC.presence_of_element_located((By.CLASS_NAME, "el-dialog"))
                    )

                    if dialog.is_displayed():
                        # 保存对话框截图
                        dialog_screenshot_path = os.path.join(IMAGE_PATH, f"{index}-{btn_index}.png")
                        driver.save_screenshot(dialog_screenshot_path)
                        print(f"[{index}] 对话框截图已保存")

                        # 查找并点击取消按钮
                        cancel_button = dialog.find_element(By.CLASS_NAME, "cancel")
                        if cancel_button.is_displayed() and cancel_button.is_enabled():
                            driver.execute_script("arguments[0].click();", cancel_button)
                            print(f"[{index}] 关闭对话框")
                            time.sleep(1)
                    else:
                        print(f"[{index}] 对话框未显示")

                except Exception as e:
                    print(f"[{index}] 等待对话框超时或处理对话框时出错:")
                    print(f"  - 错误类型: {type(e).__name__}")
                    print(f"  - 错误信息: {str(e)}")
                    # 继续处理下一个按钮
                    continue

            except Exception as e:
                print(f"[{index}] 处理按钮 {btn_index} 时出错:")
                print(f"  - 错误类型: {type(e).__name__}")
                print(f"  - 错误信息: {str(e)}")
                continue

    except Exception as e:
        print(f"[{index}] 截图过程出错:")
        print(f"  - 错误类型: {type(e).__name__}")
        print(f"  - 错误信息: {str(e)}")
# 创建一个全局线程资源管理器        
class ThreadResourceManager:
    _instance = None
    _lock = threading.Lock()
    
    def __new__(cls):
        with cls._lock:
            if cls._instance is None:
                cls._instance = super(ThreadResourceManager, cls).__new__(cls)
                cls._instance.active_threads = 0
                cls._instance.max_threads = 4  # 最大并行线程数
                cls._instance.thread_queue = queue.Queue()
                cls._instance.thread_lock = threading.Lock()
            return cls._instance
    
    def acquire_thread(self):
        """获取线程资源"""
        with self.thread_lock:
            if self.active_threads < self.max_threads:
                self.active_threads += 1
                return True
            return False
    
    def release_thread(self):
        """释放线程资源"""
        with self.thread_lock:
            if self.active_threads > 0:
                self.active_threads -= 1


def main_process(IMAGE_PATH, port, username, datetime):
    user_path = os.path.join(chrome_user_path, username, datetime)
    if not os.path.exists(user_path):
        os.makedirs(user_path)

    chrome_driver_path = os.path.join(os.path.dirname(BASE_DIR), "chromedriver", "chromedriver")
    # 配置 Selenium 使用 Chrome 浏览器
    options = webdriver.ChromeOptions()
    options.binary_location = os.path.join(os.path.dirname(BASE_DIR), "chrome", "chrome")
    options.add_argument("--disable-application-cache")
    options.add_argument("--no-sandbox")  # 禁用Linux沙箱机制
    options.add_argument(f"--user-data-dir={user_path}")  # 配置独立的用户数据目录
    options.add_argument("--disable-blink-features=AutomationControlled")  # 规避检测
    options.add_argument("--headless")  # 无头模式
    options.add_argument("--disable-gpu")  # 禁用 GPU 加速
    options.add_argument("--window-size=2580,1562")  # 设置窗口大小

    # 设置字体
    options.add_argument('--font-cache-shared-handle=0')
    options.add_argument('--lang=zh-CN')

    # 启动浏览器
    driver = webdriver.Chrome(service=Service(chrome_driver_path, port=0), options=options)

    try:
        driver.get(f"http://localhost:{port}/login")
        take_screenshot("0-0", driver, IMAGE_PATH)

        driver.get(f"http://localhost:{port}/register")
        take_screenshot("0-1", driver, IMAGE_PATH)

        # 访问 Vue 项目首页
        driver.get(f"http://localhost:{port}/")
        time.sleep(2)  # 等待页面加载
        
        # 明确点击1-1菜单项
        try:
            # 先找到菜单容器
            menu_container = driver.find_element(By.ID, "menu")
            # 然后找到第一个父菜单并点击（如果未展开）
            first_parent = menu_container.find_elements(By.CLASS_NAME, "parent-menu")[0]
            first_parent.click()
            time.sleep(0.5)  # 等待菜单展开
            
            # 找到第一个子菜单项并点击
            first_sub_item = first_parent.find_elements(By.XPATH, ".//following-sibling::ul//li")[0]
            first_sub_item.click()
            time.sleep(1)  # 等待页面响应
            
            print("成功点击了1-1菜单项")
        except Exception as e:
            print(f"点击1-1菜单项失败: {str(e)}")
            
        # 然后再截图
        take_screenshot("1-1", driver, IMAGE_PATH)


        # 收集所有需要处理的菜单项
        menu_items = []
        try:
            sidebarItems = driver.find_element(By.ID, "menu")
            parent_items = sidebarItems.find_elements(By.CLASS_NAME, "parent-menu")
            
            for index, parent in enumerate(parent_items):
                parent.click()
                time.sleep(1)  # 等待子菜单展开
                
                sub_menus = parent.find_elements(By.XPATH, ".//following-sibling::ul//li")
                for sub_index, _ in enumerate(sub_menus):
                    menu_index = f"{index+1}-{sub_index+1}"
                    if menu_index != "1-1":  # 跳过已处理的首页
                        menu_items.append((menu_index, index+1, sub_index+1))
                        
        except Exception as e:
            print(f"收集菜单项时出错: {str(e)}")
            
    finally:
        driver.quit()
    
    # 使用资源受控的线程池处理菜单项
    resource_manager = ThreadResourceManager()
    
    def process_menu_item(item_info):
        # 尝试获取线程资源
        while not resource_manager.acquire_thread():
            time.sleep(0.5)  # 等待资源释放
        
        try:
            menu_index, parent_idx, sub_idx = item_info
            
            # 为每个线程创建独立的浏览器实例
            thread_id = threading.get_ident()
            user_path = os.path.join(chrome_user_path, username, datetime, f"thread_{thread_id}")
            if not os.path.exists(user_path):
                os.makedirs(user_path)
                
            thread_options = webdriver.ChromeOptions()
            thread_options.binary_location = os.path.join(os.path.dirname(BASE_DIR), "chrome", "chrome")
            thread_options.add_argument("--disable-application-cache")
            thread_options.add_argument("--no-sandbox")
            thread_options.add_argument(f"--user-data-dir={user_path}")
            thread_options.add_argument("--disable-blink-features=AutomationControlled")
            thread_options.add_argument("--headless")
            thread_options.add_argument("--disable-gpu")
            thread_options.add_argument("--window-size=2580,1562")
            thread_options.add_argument('--font-cache-shared-handle=0')
            thread_options.add_argument('--lang=zh-CN')
            
            thread_driver = webdriver.Chrome(service=Service(chrome_driver_path), options=thread_options)
            
            try:
                # 访问首页
                thread_driver.get(f"http://localhost:{port}/")
                time.sleep(2)
                
                # 获取菜单元素
                sidebarItems = thread_driver.find_element(By.ID, "menu")
                parent_items = sidebarItems.find_elements(By.CLASS_NAME, "parent-menu")
                
                # 直接点击父菜单
                parent = parent_items[parent_idx-1]
                parent.click()
                time.sleep(1)  # 等待子菜单展开
                
                # 直接点击子菜单
                sub_menus = parent.find_elements(By.XPATH, ".//following-sibling::ul//li")
                sub_item = sub_menus[sub_idx-1]
                sub_item.click()
                
                # 截图
                take_screenshot(menu_index, thread_driver, IMAGE_PATH)
                print(f"✅ 菜单项 {menu_index} 处理完成")
                
            except Exception as e:
                print(f"❌ 处理菜单项 {menu_index} 时出错: {str(e)}")
                
            finally:
                thread_driver.quit()
                try:
                    shutil.rmtree(user_path)
                except:
                    pass
        finally:
            # 释放线程资源
            resource_manager.release_thread()
    
    # 创建线程处理所有菜单项
    threads = []
    for item in menu_items:
        thread = threading.Thread(target=process_menu_item, args=(item,))
        thread.daemon = True  # 设置为守护线程，不阻塞主程序退出
        threads.append(thread)
        thread.start()
    
    # 等待所有线程完成，但设置超时时间
    for thread in threads:
        thread.join(timeout=600)  # 最多等待10分钟
    
    print("所有菜单项处理完成或超时")

# 主函数
def main(username, datetime, IMAGE_PATH, colors):
    # 启动 Vue 项目
    vue_process, port, virtual_vue_path = start_vue_project(username, datetime, colors)
    # 截图
    main_process(IMAGE_PATH, port, username, datetime)
    # 关闭 Vue 项目
    stop_vue_project(vue_process, port, virtual_vue_path)

    arch_path = generate_architecture_diagram(username, datetime)
    if arch_path:
        print(f"架构图已生成")
    else:
        print("架构图生成失败！")


def read_txt_file(filename):
    """读取txt文件内容"""
    context = {}
    line_data = []
    flag = False
    index_local = 0
    with open(filename, "r", encoding="utf-8") as fr:
        lines = fr.readlines()
        for line in lines:
            if "*****" in line:
                flag = True
                continue
            if flag:
                if index_local == 0:
                    # 使用正则表达式清理文本
                    cleaned_line = re.sub(r'[^\w\u4e00-\u9fff]+', '', line)
                    context["name"] = cleaned_line + "页面"
                    index_local += 1
                else:
                    line_data.append(line)
    # 合并所有内容行
    content = "".join(line_data)

    # 使用正则表达式清理文本
    # 移除所有特殊字符，但保留中文、英文、数字、基本标点
    content = re.sub(r'[*{}\\/<>|]', '', content)

    # 处理多余的空白和换行
    content = re.sub(r'\n\s*\n', '\n', content)
    content = content.strip()

    context["content"] = content
    return context


def get_image_info(file):
    """获取图片信息"""
    index = file.split(".")[0]
    #print(f"\n正在查找图片 {index}.* ...")
    #print(f"在目录: {IMAGE_PATH}")
    
    try:
        image_files = os.listdir(IMAGE_PATH)
       # print(f"目录中的所有文件: {image_files}")
        
        for file_name in image_files:
            current_index = file_name.split(".")[0]
            #print(f"对比: {current_index} vs {index}")
            
            if current_index == index:
                image_path = os.path.join(IMAGE_PATH, file_name)
                if os.path.exists(image_path):
                    #print(f"找到匹配的图片: {image_path}")
                    return image_path
                else:
                    print(f"文件路径存在问题: {image_path}")
                    
        #print(f"警告: 未找到图片文件 {index}.*")
        return None
    except Exception as e:
        print(f"查找图片时出错: {str(e)}")
        return None
    

def get_sub_images(file):
    """获取带数字后缀的相关图片
    例如: 对于文件 "2-1.txt", 查找 "2-1-1.png", "2-1-2.png" 等图片
    """
    base_index = file.split(".")[0]  # 获取基础索引，如 "2-1"
    sub_images = []
    
    try:
        image_files = os.listdir(IMAGE_PATH)
        #print(f"\n查找子图片，基础索引: {base_index}")
        #print(f"在目录中的所有文件: {image_files}")
        
        for image_file in image_files:
            if image_file.startswith(f"{base_index}-"):  # 匹配前缀
                image_path = os.path.join(IMAGE_PATH, image_file)
                if os.path.exists(image_path):
                    #print(f"找到匹配的子图片: {image_path}")
                    sub_images.append(image_path)
                    
        if not sub_images:
            print(f"未找到 {base_index}-* 的子图片")
            
        return sorted(sub_images)  # 按文件名排序返回
        
    except Exception as e:
        print(f"查找子图片时出错: {str(e)}")
        return []

#生成系统架构图
def generate_system_architecture_diagram_from_menu(platform_name, save_path, menu_dict):
    dot = Digraph(comment=f'{platform_name} 系统架构图', format='png')

    # 设置全局字体与布局参数，让图更大更清晰
    dot.graph_attr.update({
        'rankdir': 'LR',          # LR：从左到右布局；也可改为 TB 自上而下
        'nodesep': '0.7',         # 节点之间的水平间距
        'ranksep': '0.7',         # 层次之间的间距
        'dpi': '150',             # 分辨率，可根据需要调整
        'fontsize': '14',         # 整体字体大小
        'fontname': 'Microsoft YaHei',
        'splines': 'ortho'        # 使用正交线条（可选）
    })
    dot.node_attr.update({
        'fontsize': '12',
        'fontname': 'Microsoft YaHei',
        'shape': 'box'
    })
    dot.edge_attr.update({
        'fontsize': '12',
        'fontname': 'Microsoft YaHei'
    })

    # 核心结构
    dot.node('FE', '前端\n(Vue)')
    dot.node('BE', '后端\n(Spring Boot)')
    dot.node('DB', '数据库\n(MySQL)', shape='cylinder')
    dot.node('EXT', '外部服务\n(预留)')

    dot.edge('FE', 'BE', label='REST API')
    dot.edge('BE', 'DB', label='数据交互')
    dot.edge('BE', 'EXT', label='预留接口')

    # 直接挂载菜单项，不做特殊处理
    for parent, children in menu_dict.items():
        # 添加父菜单节点
        dot.node(parent, parent)
        dot.edge('BE', parent)
        # 添加子菜单节点并连接
        for child in children:
            dot.node(child, child)
            dot.edge(parent, child)

    # 渲染
    diagram_base_path = os.path.join(save_path, "system_architecture")
    output_path = dot.render(diagram_base_path, cleanup=True)
    print(f"✅ 系统架构图已生成: {output_path}")
    return output_path

def generate_architecture_diagram(username, datetime_str):
    # 定义保存路径（和静态资源路径）
    save_path = os.path.join(BASE_DIR, "static", username, datetime_str)
    if not os.path.exists(save_path):
        os.makedirs(save_path)

    # 读取 config.json 以获取平台名称
    config_file = os.path.join(MEDIUM_PATH, username, datetime_str, "config.json")
    try:
        with open(config_file, "r", encoding="utf-8") as f:
            config_data = json.load(f)
        platform = config_data.get("platform", "DefaultPlatform")
    except Exception as e:
        print(f"读取配置文件失败: {e}")
        platform = "DefaultPlatform"

    # 读取 menu.json 获取侧边栏配置信息
    menu_file = os.path.join(MEDIUM_PATH, username, datetime_str, "menu.json")
    try:
        with open(menu_file, "r", encoding="utf-8") as f:
            menu_config = json.load(f)
    except Exception as e:
        print(f"读取菜单配置失败: {e}")
        menu_config = {}

    # 调用生成系统架构图的函数
    try:
        architecture_path = generate_system_architecture_diagram_from_menu(platform, save_path, menu_config)
        print(f"架构图生成成功，路径：{architecture_path}")
        return architecture_path
    except Exception as e:
        print(f"生成系统架构图失败: {e}")
        return None


def generate_word_template(title, user, time, TXT_PATH):
    doc = Document(os.path.join(BASE_DIR, "medium", "template.docx"))
    version_str = "V1.0"
    main_info = {
        "title": title,
        "table": [{"version": version_str, "date": time, "name": user, "info": "初始版本"}]
    }
    # 去掉“环境描述”，将“开发目的”修改为“系统简介”
    config = [
        {"name": "系统简介", "content": ""},
        {"name": "功能描述", "content": ""}
    ]
    main_content = {"title": "使用说明", "subsection": []}

    # 设置页眉
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.clear()
    run = header_paragraph.add_run(f"{title} {version_str}")
    run.font.name = "宋体"
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 替换 {{ title }} 为标题 + 操作手册
    for paragraph in doc.paragraphs:
        if "{{ title }}" in paragraph.text:
            paragraph.clear()
            # 添加标题（标题部分采用一号字）
            run1 = paragraph.add_run(f"{title} {version_str}")
            run1.bold = True
            run1.font.name = "黑体"
            run1.font.size = Pt(42)  # 调整为一号字

            # 添加换行（两个换行符，中间空一行）
            paragraph.add_run("\n")
            paragraph.add_run("\n")

            # 添加“操作手册”文字
            run2 = paragraph.add_run("操作手册")
            run2.bold = True
            run2.font.name = "黑体"
            run2.font.size = Pt(40)

            # 设置居中及行间距为双倍行距
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.paragraph_format.line_spacing = 2
            break

    # 生成配置信息
    config_content = generate_config(title, user, time)

    config[0]["content"] = config_content.get("系统简介", "")
    config[1]["content"] = config_content.get("功能描述", "")

    # 作者信息写入表格
    table = doc.tables[0]
    for i, row_data in enumerate(main_info["table"]):
        table.cell(i + 1, 0).text = row_data["version"]
        table.cell(i + 1, 1).text = str(row_data["date"])
        table.cell(i + 1, 2).text = row_data["name"]
        table.cell(i + 1, 3).text = row_data["info"]

    # 平台配置信息
    for section in config:
        doc.add_page_break()
        doc.add_heading(section["name"], level=1)
        if "subsection" in section.keys():
            for subsection in section["subsection"]:
                doc.add_heading(subsection["name"], level=2)
                doc.add_paragraph(subsection["content"])
        else:
            if type(section["content"]) == str:
                doc.add_paragraph(section["content"])
            else:
                # 添加带序号的段落
                doc = add_multi_level(doc, section["content"])

    # 添加第五部分：系统设计（架构图）
    doc.add_page_break()
    doc.add_heading("系统设计", level=1)
    doc.add_heading("系统架构图", level=2)
    try:
        system_img_path = os.path.join(BASE_DIR, "static", user, time, "system_architecture.png")
        if os.path.exists(system_img_path):
            para = doc.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.add_run().add_picture(system_img_path, height=Cm(21.33))
        else:
            print("⚠️ 未找到系统架构图:", system_img_path)
    except Exception as e:
        print(f"插入系统架构图失败: {e}")

    # 读取 TXT 并插入对应图片
    files = natsorted(os.listdir(TXT_PATH))
    for file in files:
        filename = os.path.join(TXT_PATH, file)
        context = read_txt_file(filename)

        # 获取主图片
        image = get_image_info(file)
        if image:
            context["image"] = image

            # 获取子图片
            sub_images = get_sub_images(file)
            if sub_images:
                context["sub_images"] = sub_images

            main_content["subsection"].append(context)
        else:
            print(f"跳过 {file} 的处理，因为没有找到对应的图片")

    doc.add_page_break()
    doc.add_heading(main_content["title"], level=1)

    for section in main_content["subsection"]:
        # 添加标题
        doc.add_heading(section["name"], level=2)

        # 1. 添加主图片 - 优先使用标注版本
        if "image" in section and section["image"]:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 检查是否存在标注版本
            image_path = section["image"]
            name, ext = os.path.splitext(image_path)
            annotated_path = f"{name}_annotated{ext}"

            try:
                if os.path.exists(annotated_path):
                    paragraph.add_run().add_picture(annotated_path, width=Inches(5.5))
                    print(f"使用标注版图片: {annotated_path}")
                else:
                    paragraph.add_run().add_picture(image_path, width=Inches(5.5))
                    print(f"使用原始图片: {image_path}")
            except Exception as e:
                print(f"添加图片时出错: {str(e)}")

        # 2. 添加内容描述
        doc.add_paragraph(section["content"])

        # 3. 添加子图片 - 同样优先使用标注版本
        if "sub_images" in section:
            for sub_image in section["sub_images"]:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # 检查是否存在标注版本
                name, ext = os.path.splitext(sub_image)
                annotated_sub_path = f"{name}_annotated{ext}"

                try:
                    if os.path.exists(annotated_sub_path):
                        paragraph.add_run().add_picture(annotated_sub_path, width=Inches(5.5))
                        print(f"使用标注版子图片: {annotated_sub_path}")
                    else:
                        paragraph.add_run().add_picture(sub_image, width=Inches(5.5))
                        print(f"使用原始子图片: {sub_image}")
                except Exception as e:
                    print(f"添加子图片时出错: {str(e)}")

    add_pager_header(doc, f"{platform} {version_str}")

    doc_save_path = os.path.join(BASE_DIR, "static", user, time, "template_manual.docx")
    doc.save(doc_save_path)
    print('Word 文档生成完毕')


if __name__ == "__main__":
    platform = sys.argv[1]
    username = sys.argv[2]
    datetime = sys.argv[3]
    colors = sys.argv[4]

    final_path = os.path.join(BASE_DIR, "static", username, datetime)
    if not os.path.exists(final_path):
        os.makedirs(final_path)

    TXT_PATH = os.path.join(BASE_DIR, "Introduction", username, datetime)
    IMAGE_PATH = os.path.join(BASE_DIR, "screenshot", username, datetime)
    if not os.path.exists(IMAGE_PATH):
        os.makedirs(IMAGE_PATH)
    if not os.path.exists(TXT_PATH):
        os.makedirs(TXT_PATH)

    main(username, datetime, IMAGE_PATH, colors)
    generate_word_template(platform, username, datetime, TXT_PATH)

