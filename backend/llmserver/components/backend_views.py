from openai import OpenAI
import os
import json
import random
from docx import Document
import re
import concurrent.futures
import sys
from pdf_views import main_pdf
from docx.shared import Pt
from docx.oxml.ns import qn
import time
from tenacity import retry, stop_after_attempt, wait_exponential
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.append(BASE_DIR)

from backend import settings

client = OpenAI(
    api_key=settings.BACKEND_LLM_KEY[0],
    base_url=settings.BACKEND_LLM_URL[0],
)
@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=20))
def get_model_register_response(messages, model="qwq-plus-latest"):
    try:
        print(f"[DEBUG] 发送请求到模型: {model}")
        # 使用流式模式调用模型
        response_stream = client.chat.completions.create(
            model=model,
            messages=messages,
            stream=True  # 启用流式响应
        )
        
        # 从流中收集完整响应
        full_response = ""
        for chunk in response_stream:
            if chunk.choices and hasattr(chunk.choices[0], 'delta') and hasattr(chunk.choices[0].delta, 'content'):
                content = chunk.choices[0].delta.content
                if content is not None:
                    full_response += content
        
        print(f"[DEBUG] 模型响应完成，长度: {len(full_response)}")
        return full_response.strip()
    except Exception as e:
        print(f"[ERROR] 模型调用失败: {e}")
        raise e

def get_model_response(user_message):
    message = [{"role": "user", "content": user_message}]
    try:
        completion = client.chat.completions.create(
            model="qwen-plus",
            messages=message
        )
        return json.loads(completion.model_dump_json())["choices"][0]["message"]["content"]
    except Exception as e:
        raise e

# 定义调用大模型的函数，用于代码生成
def get_response(messages):
    try:
        completion = client.chat.completions.create(
            model="qwen2.5-coder-32b-instruct",
            messages=messages
        )
        return completion
    except Exception as e:
        raise e


# 保存模型输出内容到txt文件
def save_to_txt(content, filename):
    try:
        with open(filename, "w", encoding="utf-8") as file:
            file.write(content)
        print(f"内容已保存到 {filename}")
    except Exception as e:
        print(f"保存到txt文件时出错: {e}")

# 保存模型输出内容到文档
def save_to_word(content, filename):
    try:
        doc = Document()
        doc.add_paragraph(content)
        doc.save(filename)
        print(f"内容已保存到 Word 文档: {filename}")
    except Exception as e:
        print(f"保存到Word文档时出错: {e}")

# 读取指定路径的 JSON 文件
def read_json_file(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            data = json.load(file)
            return data
    except Exception as e:
        print(f"读取文件时出错: {e}")
        return None


# 读取 JSON 文件内容并调用大模型理解
def analyze_file(file_path, txt_file_path):
    file_data = read_json_file(file_path)

    if file_data is not None:
        # 将文件内容转换为字符串并传递给大模型
        file_content = json.dumps(file_data, ensure_ascii=False)
        question = f"请理解以下JSON文件内容并回答：{file_content}"

        # 调用大模型获取回答
        answer = get_model_response(question)

        # print("最终答案：")
        # print(answer)

        # 将思考过程和答案保存到txt文件
        save_to_txt(f"最终答案：\n{answer}", txt_file_path)

        # 返回答案以供后续使用
        return answer
    else:
        print("无法读取文件，无法继续分析。")
        return None


# 第一次提问：扩展软件信息
def first_prompt(software_name, programming_language, txt_file_path, final_path):
    messages = [{
        "role": "system",
        "content": "You are a helpful assistant that expands software descriptions based on input."
    }]
    # # 获取编程语言的输入
    # programming_language = input("请输入软件的编程语言：")
    # 扩展软件描述的模板
    user_input = f"请根据以下框架信息扩展对{software_name}这一软件著作权的软件描述\n"
    user_input += f"""
    扩展软件描述的模板如下：
    【软件全称】{software_name}
    【版本号】
    【软件分类】
    【开发的硬件环境】
    【运行的硬件环境】
    【开发该软件的操作系统】
    【软件开发环境/开发工具】
    【该软件的运行平台/操作系统】
    【软件运行支撑环境/支持软件】
    【编程语言】{programming_language}  # 使用用户输入的编程语言
    【源程序量】
    【开发目的】
    【面向领域/行业】
    【软件的主要功能】
    【软件的技术特点】
    请注意:【源程序量】后面仅需要填行数,行数是生成的随机数，范围为20000-60000行。其中各位均不为0，【版本号】写死为v1.0
    
    """

    messages.append({"role": "user", "content": user_input})

    # 获取模型返回的完整内容
    assistant_output = get_response(messages).choices[0].message.content

    # print(f"第一次提问的输出：\n{assistant_output}")

    # 保存到指定txt文件
    save_to_txt(assistant_output, txt_file_path)
    save_to_txt(assistant_output, final_path)

    return assistant_output


# 第二次提问：生成代码框架，包含模块分析
def second_prompt(programming_language, expanded_description, txt_file_path):
    messages = [{
        "role": "system",
        "content": f"You are a helpful assistant that generates {programming_language} code framework based on software descriptions, and includes detailed descriptions and comments for each module."
    }]
    user_input = f"根据以下软件描述，生成{programming_language}代码框架，并为每个模块加入详细描述和注释：\n{expanded_description}\n"
    user_input += """
    并请详细描述代码框架的各个部分：
    - 前端部分：包括用户界面设计、交互逻辑、状态管理、API调用等，每个模块需要有详细的注释和描述。
    - 后端部分：包括用户认证模块、数据存储模块、错误处理模块等，每个模块需要有详细的注释和描述。
    - 数据库部分：包括多个数据库表设计、存储过程、视图等，每个部分需要详细描述。
    - API部分：设计多个RESTful API接口，并加入权限控制、认证、限流等功能。
    - 业务部分：包括用户管理、订单处理等复杂业务逻辑模块，每个模块需要有详细的描述。
    - 安全部分：包括数据加密、认证与授权、日志记录等每个部分需要有详细说明。
    """
    messages.append({"role": "user", "content": user_input})

    assistant_output = get_response(messages).choices[0].message.content
    # print(f"第二次提问的输出：\n{assistant_output}")

    # 保存到指定txt文件
    save_to_txt(assistant_output, txt_file_path)

    return assistant_output


# 修改后的代码生成函数
def generate_code_with_details(programming_language, layer, code_framework):
    messages = [{
        "role": "system",
        "content": f"You are a helpful assistant that generates {programming_language} code with detailed comments and explanations for each module. Please only return the code and remove any extra explanations or descriptions."
    }]
    user_input = f"理解下面描述中关于{layer}的部分，并为每个模块生成详细注释的{programming_language}代码, 且代码注释必须为中文，不可以使用英文注释。不包含任何额外的描述，且代码内容不应包含其他格式的解释或说明，不少于500行：\n{code_framework}\n"
    messages.append({"role": "user", "content": user_input})

    assistant_output = get_response(messages).choices[0].message.content

    # 过滤掉不需要的内容：删除以“Below is a detailed...”开头的描述性文字、代码块标记、注释等
    # 删除以 ``` 开头的行和其余内容
    code_only = re.sub(r"(?i)^(below is a detailed.*|python\s*|code follows.*|the code is).*", "", assistant_output)
    code_only = re.sub(r"^```.*", "", code_only)  # 删除以 ``` 开头的行

    # 删除空行，确保代码紧凑
    code_only = "\n".join([line for line in code_only.split("\n") if line.strip() != ""])

    # print(f"{layer}层生成的代码：\n{code_only}")
    return code_only


# 细化每个模块的代码生成，修改生成“用户表设计”模块的部分
def generate_submodules_for_layer(programming_language, layer_name, code_framework, json_answer=None):
    submodules = {
        "后端": ["用户认证模块", "数据存储模块", "错误处理模块", "日志记录模块", "配置管理模块"],
        "数据库": ["用户表设计", "存储过程", "视图", "索引优化", "数据备份机制"],
        "API": ["用户API接口", "订单API接口", "认证接口", "限流接口", "第三方API对接"],
        "业务": ["用户管理", "订单处理", "支付处理", "库存管理", "通知系统"],
        "安全": ["数据加密", "权限验证", "访问控制", "输入输出校验", "登录日志追踪"],
    }

    all_code = []
    for submodule in submodules[layer_name]:
        submodule_description = f"{layer_name}层 - {submodule}"

        # 如果是“用户表设计”模块，将json_answer加入到输入
        if layer_name == "数据库" and submodule == "用户表设计" and json_answer:
            submodule_code = generate_code_with_details(programming_language, submodule_description, code_framework + "\n" + json_answer)
        else:
            submodule_code = generate_code_with_details(programming_language, submodule_description, code_framework)

        all_code.append(submodule_code)

    return "\n".join(all_code)


# 并发生成所有层代码
def generate_all_layers_concurrently(programming_language, code_framework, json_answer=None):
    layer_names = [ "后端", "数据库", "API", "业务", "安全"]

    with concurrent.futures.ThreadPoolExecutor() as executor:
        future_to_layer = {
            executor.submit(generate_submodules_for_layer, programming_language, layer, code_framework,
                            json_answer if layer == "数据库" else None): layer
            for layer in layer_names
        }

        results = {}
        for future in concurrent.futures.as_completed(future_to_layer):
            layer_name = future_to_layer[future]
            try:
                results[layer_name] = future.result()
                print(f"{layer_name} 代码生成完成！")
            except Exception as e:
                print(f" {layer_name} 代码生成失败: {e}")

    return results


# 合并多个Word文档为一个
def merge_word_documents(file_paths, output_path, software_name):
    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    merged_doc = Document()

    # 添加页眉
    section = merged_doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = software_name
    header_paragraph.alignment = 1  # 居中

    run = header_paragraph.runs[0]
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    run.font.size = Pt(10.5)

    # 遍历文件合并内容
    for file_path in file_paths:
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text = para.text

                # 依次去除代码块标识符
                text = text.replace("```\n", "").replace("```", "")

                # 依次去除各种编程语言名称（大小写都考虑）
                languages = [
                    "python", "Python", "java", "Java", "Php","php"
                ]

                for lang in languages:
                    text = text.replace(lang, "")

                new_para = merged_doc.add_paragraph(text)
                for run in new_para.runs:
                    run.font.name = "SimSun"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
                    run.font.size = Pt(10.5)

            merged_doc.add_page_break()  # 添加分页符
            print(f"成功合并文档: {file_path}")
        except Exception as e:
            print(f"合并文档时出错: {file_path}, 错误: {e}")

    # 等待2秒并保存
    time.sleep(2)
    try:
        merged_doc.save(output_path)
        print(f"所有文档已合并并保存到: {output_path}")
    except Exception as e:
        print(f"保存合并文档时出错: {e}")

def extract_technical_features_from_file(username, datetime):
    """从function.txt文件中获取技术特点"""
    try:
        # 构建function.txt路径
        function_file_path = os.path.join(BASE_DIR, "medium", username, datetime, "function.txt")
        print(f"尝试读取功能文件: {function_file_path}")
        while True:
            if os.path.exists(function_file_path):
            # 如果存在，直接读取其内容作为技术特点
                with open(function_file_path, 'r', encoding='utf-8') as f:
                    tech_features = f.read().strip()
                    print(f"✅ 从function.txt成功读取技术特点，长度: {len(tech_features)}")
                    break
        return tech_features
    except Exception as e:
        print(f"❌ 读取文件出错: {e}")
        return ""
def generate_source_lines():
    while True:
        num = random.randint(20000, 25000)
        if '0' not in str(num):
            return num
# === 生成标准字段（通用字段9项）
def generate_fields_via_model():
    prompt = f"""
    请根据以下框架信息扩展对{software_name}这一软件著作权的软件描述\n
请为以下软件字段生成内容。要求：
1. 每项内容直接接在字段标签后（同一行）；
2. 内容风格为简洁罗列；
3. 每项独占一行，示例如下：

【开发的硬件环境】CPU: Intel i7-10700K 主频:3.8GHz 内存:32GB SSD:512GB
【运行的硬件环境】CPU: Xeon E5-2640 内存:16GB 存储:1.2T*24块(SAS)
【开发该软件的操作系统】Windows Server 2016、Ubuntu 20.04
【软件开发环境或者开发工具】PyCharm、VS Code、Git、MySQL Workbench
【软件的运行平台或操作系统】Windows 10、Android 11、Linux CentOS
【软件运行支撑环境或支持软件】Python 3.10、Redis、MySQL、Chrome
【开发目的】
【面向领域/行业】

⚠️只生成以上字段，禁止生成多余字段。
"""
    messages = [{'role': 'user', 'content': prompt}]
    return get_model_register_response(messages)
# === 生成主要功能描述（简洁版，不判定长度）===
def generate_main_function(software_name,tech_features):
    prompt = f"""
请围绕一款名为“{software_name}”的软件，根据已有的软件功能总结生成一段中文描述说明其主要功能，务必注意内容长度必须大于100个字符，同时小于200个字符。\n
{tech_features}
"""
    messages = [{'role': 'user', 'content': prompt}]
    response = client.chat.completions.create(
        model="qwen2.5-coder-32b-instruct",
        messages=messages
    )
    return response.choices[0].message.content.strip()
# === 总结技术特点（20~100字）
def summarize_technical_features(raw_features):
    prompt = f"""
请将以下技术特点内容压缩为20到100个汉字之间的简明描述，保持主要含义，风格专业：
{raw_features}
只返回压缩后的内容。
"""
    messages = [{'role': 'user', 'content': prompt}]
    response = client.chat.completions.create(
        model="qwen2.5-coder-32b-instruct",
        messages=messages
    )
    return response.choices[0].message.content.strip()
# === 总结技术特点（20~100字）
def summarize_technical_features(raw_features):
    prompt = f"""
请将以下技术特点内容压缩为20到100个汉字之间的简明描述，保持主要含义，风格专业：
{raw_features}
只返回压缩后的内容。
回复内容中除了压缩的内容，不要包括任何其余解释
"""
    messages = [{'role': 'user', 'content': prompt}]
    return get_model_register_response(messages)

# === 根据技术特点内容判断最合适的选项
def choose_tech_option(summary):
    options = ["APP", "游戏软件", "人工智能软件", "金融软件", "大数据软件", "云计算软件", "信息安全软件", "小程序", "物联网", "智慧城市软件"]
    prompt = f"""
根据以下软件技术特点内容，判断最适合归类的类型，并从以下选项中选择一个最匹配的：
{options}

技术特点描述如下：
{summary}

请返回最合适的一个类别名称（仅输出选项本身，不加任何解释）。
"""
    messages = [{'role': 'user', 'content': prompt}]
    return get_model_register_response(messages)
        
# === 拼接最终内容并保存
def generate_and_save_txt(output_path):
    global software_name, programming_language, tech_features
    
    version = "V1.0"
    category = "应用软件"
    code_lines = generate_source_lines()

    # 生成所有部分
    model_fields = generate_fields_via_model().splitlines()
    main_function = generate_main_function(software_name,tech_features)
    summarized_tech = summarize_technical_features(tech_features)
    tech_option = choose_tech_option(summarized_tech)

    # 插入语言/行数
    insert_after = "【软件运行支撑环境或支持软件】"
    insert_index = next((i for i, line in enumerate(model_fields) if line.startswith(insert_after)), -1)
    if insert_index != -1:
        model_fields.insert(insert_index + 1, f"【编程语言】{programming_language}")
        model_fields.insert(insert_index + 2, f"【源程序量】{code_lines}")
    else:
        model_fields.append(f"【编程语言】{programming_language}")
        model_fields.append(f"【源程序量】{code_lines}")

    # 拼接尾部字段
    model_fields.append(f"【软件的主要功能】{main_function}")
    model_fields.append(f"【技术特点】{summarized_tech}")
    model_fields.append(f"【软件的技术特点选项】{tech_option}")

    # 最终合成
    final_text = f"""【软件名称】{software_name}
【版本号】{version}
【软件分类】{category}
""" + "\n".join(model_fields)

    try:
        # 始终保存为TXT格式文件
        txt_output_path = output_path
        if output_path.endswith('.docx'):
            txt_output_path = output_path.replace('.docx', '.txt')
        
        # 保存为文本文件
        with open(txt_output_path, "w", encoding="utf-8") as f:
            f.write(final_text)
        print(f"✅ 成功保存文本文件至：{txt_output_path}")
    except Exception as e:
        print(f"❌ 写入失败：{e}")



# 运行主要流程
if __name__ == "__main__":
    platform = sys.argv[1]
    language = sys.argv[2]
    username = sys.argv[3]
    datetime = sys.argv[4]
    software_name = platform
    programming_language = language
    output_path=os.path.join(BASE_DIR, "static", username, datetime, "软著注册表.txt")
    medium_path = os.path.join(BASE_DIR, "medium", username, datetime)
    final_path = os.path.join(BASE_DIR, "static", username, datetime)
    if not os.path.exists(medium_path):
        os.makedirs(medium_path)
    if not os.path.exists(final_path):
        os.makedirs(final_path)
    # 定义保存代码的路径
    output_path_1 = medium_path  # 保存各层的代码文档
    output_path_2 = final_path  # 保存合并后的总代码文档

# 调用函数，传入文件路径
    file_path = os.path.join(medium_path, "menu.json")
    json_answer = analyze_file(file_path, os.path.join(medium_path, "json_analysis.txt"))
    # 第1次提问:扩充软著描述，第2次提问：软著实现整体框架
  
    #expanded_description = first_prompt(platform, language, os.path.join(medium_path, "expanded_description.txt"),
    #                                    os.path.join(final_path, "expanded_description.txt"))  # 第一次提问，获得对软著名的扩充描述，并保存到txt
    code_framework = second_prompt(language, json_answer,os.path.join(medium_path, "code_framework.txt"))  # 第二次提问，获得对软著实现的框架性语言，并保存到txt

    # 并发生成代码
    layer_codes = generate_all_layers_concurrently(language, code_framework, json_answer)

    # 保存各层代码到 Word 文件
    file_paths = []  # 存储各层代码word文档的地址，方便后续合并
    
    # 持续检查前端代码文件是否存在，直到找到为止
    frontend_code_path = os.path.join(medium_path, "前端_code.docx")
    retry_interval = 5  # 每次等待5秒
    attempt = 0
    
    print("开始检查前端代码文件...")
    while True:
        if os.path.exists(frontend_code_path):
            file_paths.append(frontend_code_path)
            print(f"检测到前端代码文件，已添加到合并列表: {frontend_code_path}")
            break
    
    
    # 添加后端层代码文件
    for layer, code in layer_codes.items():  # 获取键值对，层名-代码
        file_path = os.path.join(output_path_1, f"{layer}_code.docx")
        save_to_word(code, file_path)
        file_paths.append(file_path)

    # 合并所有代码文件
    merged_file_path = os.path.join(output_path_2, "merged_code.docx")
    merge_word_documents(file_paths, merged_file_path, platform)
    main_pdf(username, datetime)
    #生成软著注册表
    #读取function.txt文件
    tech_features = extract_technical_features_from_file(username, datetime)
    #生成软著注册表
    print(f"技术特点: {tech_features}")
    print(f"生成软著注册表: 软件名称={software_name}, 编程语言={programming_language}")
    generate_and_save_txt(output_path)
