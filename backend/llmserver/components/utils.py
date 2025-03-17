from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import sys
from openai import OpenAI
import json
from tenacity import retry, stop_after_attempt, wait_exponential
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.append(BASE_DIR)

from backend import settings

client = OpenAI(
    api_key=settings.LLM_API_KEY,
    base_url=settings.LLM_BASE_URL
)



# 设置带序号的段落
def add_multi_level(doc, content):
    for i, item in enumerate(content.keys(), start=1):
        # 添加主段落
        doc.add_paragraph(f"({i}) {item}")

        # 添加子段落
        for j, sub_item in enumerate(content[item].keys(), start=1):
            sub_paragraph = doc.add_paragraph(f"({i}-{j}) {sub_item}: {content[item][sub_item]}")
            # 设置子段落缩进
            sub_paragraph.paragraph_format.left_indent = Inches(0.5)  # 缩进 0.5 英寸

    return doc

@retry(stop=stop_after_attempt(10), wait=wait_exponential(multiplier=1, min=2, max=20))
def api_message(introduction, platform, subtitle):
    try:
        MESSAGE = [{"role": "system", "content": "You are a helpful programmer and product manager"}]
        question = f"""
            1.现在有一个叫{platform}的后台管理系统.下面有一个名叫{subtitle}的页面.
            2.请理解该页面的设计说明文字: "{introduction}"并对该页面布局有大致的掌握.
            """
        query = {"role": "user", "content": question}
        MESSAGE.append(query)
        completion = client.chat.completions.create(
            model="deepseek-r1",
            messages=MESSAGE
        )
        question = f"""
        1.根据该页面的布局生成面向用户的使用说明,介绍该页面具体的功能以及操作指南.
        2.生成的内容每一句之间确保有正确的中文标点符号.
        3.生成的文字内容格式上无需换行.
        4.不需要提供功能说明之外的其他布局信息以及代码编写信息,例如组件间隔,基于什么框架等.
        """
        query = {"role": "system", "content": question}
        MESSAGE.append(query)
        completion = client.chat.completions.create(
            model="deepseek-r1",
            messages=MESSAGE
        )
        data_dict = json.loads(completion.model_dump_json())
        content = data_dict["choices"][0]["message"]["content"]
        return content
    except Exception as e:
        raise e


# 添加每个页面的文字说明
def add_manual(section, platform, i, content_dict):
    content = api_message(section["content"], platform, section["name"])
    print(f"{section['name']}页面生成完毕")
    content_dict[str(i)] = content


# 添加页眉和页码
def add_pager_header(doc, platform):
    section = doc.sections[0]
    header = section.header

    # 清除页眉内容（通过删除所有段落）
    for paragraph in header.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    header_paragraph = header.add_paragraph()
    header_paragraph.add_run(platform)

    p_pr = header_paragraph._element.get_or_add_pPr()
    bottom_border = OxmlElement('w:pBdr')
    bottom_border_line = OxmlElement('w:bottom')
    bottom_border_line.set(qn('w:val'), 'single')  # 设置边框样式为单线
    bottom_border_line.set(qn('w:sz'), '6')  # 设置边框宽度（单位：1/8磅）
    bottom_border_line.set(qn('w:space'), '1')  # 设置边框与文本的间距
    bottom_border_line.set(qn('w:color'), '000000')  # 设置边框颜色（黑色）
    bottom_border.append(bottom_border_line)
    p_pr.append(bottom_border)

    header_paragraph.alignment = 0

    header_paragraph.add_run().add_tab()

    # 插入页码字段
    run = header_paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

    # 设置段落制表位以实现页码右对齐
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9000')
    tabs.append(tab)
    p_pr = header_paragraph._element.get_or_add_pPr()
    p_pr.append(tabs)

