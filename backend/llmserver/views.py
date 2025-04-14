from django.views.decorators.http import require_http_methods
from django.http import JsonResponse, FileResponse
from openai import OpenAI
import os
import json
import threading
import random
import shutil
from tenacity import retry, stop_after_attempt, wait_exponential
import shutil
from http import HTTPStatus
from urllib.parse import urlparse, unquote
from pathlib import PurePosixPath
import requests
from dashscope import ImageSynthesis
from urllib.parse import quote
import os
from docx import Document
import re
import subprocess
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "backend.backend.settings")
from django.conf import settings
from .models import *

from django.db.models.fields import DateTimeField
from django.db.models.fields.related import ManyToManyField

RANDOM_CHOICE = random.randint(1, 2)

# 初始化 OpenAI 客户端，增加 timeout 参数
client = OpenAI(
    api_key=settings.LLM_API_KEY,
    base_url=settings.LLM_BASE_URL,
    timeout=300
)
api_keynum = settings.LLM_API_KEY
MAX_THREADS = 3
WEB_URL = "http://121.196.229.117:8000/static"
IMG_URL = "http://121.196.229.117/static"
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MEDIUM_PATH = os.path.join(BASE_DIR, "medium")

# def build_static_url(username, datetime_str, filename):
#     return f"{settings.STATIC_URL}{quote(username)}/{quote(datetime_str)}/{filename}"


def read_json(username, datetime_str):
    config_path = os.path.join(MEDIUM_PATH, username, datetime_str, "config.json")
    with open(config_path, "r", encoding="utf-8") as fr:
        data = json.load(fr)
    return data


def delete_directory(folder_path):
    if os.path.exists(folder_path):
        # 删除文件夹及其所有内容
        shutil.rmtree(folder_path)
    else:
        raise Exception("文件夹不存在")


def get_response(message):
    """
    使用流式请求模型，并将分块返回的内容合并为完整回复
    """
    reasoning_content = ""  # 定义完整思考过程
    full_response = ""     # 修改这里，使用 content 而不是 answer_content
    is_answering = False
    try:
        print("[DEBUG] 模型请求发送中...")
        # 采用流式返回，model 使用 qwq-32b（可根据实际情况调整）
        completion = client.chat.completions.create(
            model="qwq-plus-latest",
            messages=message,
            stream=True
        )
        #print("\n" + "=" * 20 + "思考过程" + "=" * 20 + "\n")
        for chunk in completion:
            if not chunk.choices:
                print("\nUsage:")
                #print(chunk.usage)
            else:
                delta = chunk.choices[0].delta
                if hasattr(delta, 'reasoning_content') and delta.reasoning_content != None:
                    #print(delta.reasoning_content, end='', flush=True)
                    reasoning_content += delta.reasoning_content
                else:
                    if delta.content != "" and is_answering is False:
                        #print("\n" + "=" * 20 + "完整回复" + "=" * 20 + "\n")
                        is_answering = True
                    #print(delta.content, end='', flush=True)
                    full_response += delta.content
        return full_response
    except Exception as e:
        print(f"[ERROR] 模型调用失败: {e}")
        raise e


@retry(stop=stop_after_attempt(10), wait=wait_exponential(multiplier=1, min=2, max=20))
def api_call(message):
    return get_response(message)


def to_dict(self, fields=None, exclude=None):
    data = {}
    for f in self._meta.concrete_fields + self._meta.many_to_many:
        value = f.value_from_object(self)

        if fields and f.name not in fields:
            continue

        if exclude and f.name in exclude:
            continue

        if isinstance(f, ManyToManyField):
            value = [i.id for i in value] if self.pk else None

        if isinstance(f, DateTimeField):
            value = value.strftime("%Y-%m-%d %H:%M:%S") if value else None

        data[f.name] = value

    return data


def run_another_script_webfront(arg1, arg2, arg3, arg4):
    global MAX_THREADS
    import subprocess
    if isinstance(arg4, list):
        arg4 = ",".join(arg4)
    subprocess.run(
        ["python3", os.path.join(BASE_DIR, "llmserver", "components", "webfront_screenshot.py"), arg1, arg2, arg3, arg4]
    )
    MAX_THREADS += 1


def run_another_script_backend(arg1, arg2, arg3, arg4):
    # 使用 subprocess.run 运行另一个 Python 文件并传递参数
    import subprocess
    subprocess.run(
        ["python3", os.path.join(BASE_DIR, "llmserver", "components", "backend_views.py"), arg1, arg2, arg3, arg4]
    )


@require_http_methods(["GET"])
def check_thread_pool_available(request):
    response = {"code": 0, "message": "success"}
    response["status"] = 1 if MAX_THREADS > 0 else 0
    return JsonResponse(response)


@require_http_methods(["POST"])
def startProgram(request):
    global MAX_THREADS
    MAX_THREADS -= 1
    response = {"code": 0, "message": "success"}
    data = json.loads(request.body)
    person_id = data.get("id")
    platform = data.get("platform")
    language = data.get("language")
    datetime_str = data.get("time")
    username = data.get("username")
    colors = data.get("color")

    code_path = f"{WEB_URL}/{username}/{datetime_str}/ultimate_file.pdf"
    word_file = f"{WEB_URL}/{username}/{datetime_str}/template_manual.docx"
    introduce_file = f"{WEB_URL}/{username}/{datetime_str}/expanded_description.txt"

    # print("[DEBUG] 生成的文件访问路径：")
    # print("PDF下载地址：", code_path)
    # print("Word文档地址：", word_file)
    # print("介绍文档地址：", introduce_file)


    medium_url = os.path.join(MEDIUM_PATH, username, datetime_str) # 用户存放某用户在某时间下创建的项目的具体信息
    if not os.path.exists(medium_url):
        os.makedirs(medium_url)
    user_dict = {
        "platform": platform,
        "username": username,
        "datetime": datetime_str,
        "language": language
    }
    with open(os.path.join(medium_url, "config.json"), "w", encoding="utf-8") as f:
        json.dump(user_dict, f, indent=4)

    record = UserRecord.objects.create(
        name=platform,
        time=datetime_str,
        language=language,
        user=UserProfile(id=person_id),
        pdf_download=word_file,
        code_download=code_path,
        introduce_download=introduce_file
    )
    record.save()
    # 生成侧边栏信息
    MESSAGE = [{"role": "system", "content": "You are a helpful programmer and product manager"}]
    question = f"""
    设计一个{platform}的侧边栏, 有如下要求:
    1. 侧边栏需包含至少5个目录，每个目录名称应与系统业务逻辑密切相关。
    2. 返回结果应为一个Object，每个Object的key为父目录，对应的value为一个List，包含所有子目录名称。
    3. 回答时只需返回上述Object，无需任何解释说明，且返回内容必须以Object的"{{"开始。
    """
    MESSAGE.append({"role": "user", "content": question})
    try:
        full_reply = api_call(MESSAGE)
        menu = json.loads(full_reply)
    except Exception as e:
        print(f"侧边栏生成失败: {e}")
        menu = {"主菜单": ["数据统计", "消息通知"], "系统管理": ["用户管理", "权限设置"]}
    
    MENU_CONFIG = menu
    json_file = os.path.join(MEDIUM_PATH, username, datetime_str, "menu.json")
    with open(json_file, "w", encoding="utf-8") as f:
        json.dump(MENU_CONFIG, f)
    
    # 启动后端代码生成线程
    thread_backend = threading.Thread(
        target=run_another_script_backend,
        args=(platform, language, username, datetime_str)
    )
    thread_backend.start()

    thread1 = threading.Thread(
        target=run_another_script_webfront,
        args=(platform, username, datetime_str, colors)
    )
    thread1.start()
    return JsonResponse(response)


@require_http_methods(["POST"])
def login(request):
    response = {"code": 2000, "message": "success"}
    data = json.loads(request.body)
    username = data.get("username")
    password = data.get("password")
    user_count = UserProfile.objects.filter(username=username).count()
    if user_count == 0:
        response["code"] = 1
        response["message"] = "User not found"
        return JsonResponse(response)
    else:
        if not UserProfile.objects.filter(username=username, password=password).count():
            response["code"] = 2
            response["message"] = "Password Wrong"
            return JsonResponse(response)
        else:
            user = UserProfile.objects.get(username=username, password=password)
            data = {
                "username": username,
                "roles": [user.role],
                "id": user.id
            }
            response["data"] = data
            response["code"] = 0
            return JsonResponse(response)


@require_http_methods(["GET"])
def getuserinfo(request):
    response = {"code": 2000, "message": "success"}
    user_id = request.GET.get("user_id")
    user = UserProfile.objects.get(id=user_id)
    response["data"] = to_dict(user)
    return JsonResponse(response)


@require_http_methods(["POST"])
def editUserInfo(request):
    response = {"code": 2000, "message": "success"}
    data = json.loads(request.body)
    username = data.get("username")
    password = data.get("password")
    email = data.get("email")
    phone = data.get("phone_number")
    user_id = data.get("id")
    user = UserProfile.objects.get(id=user_id)
    user.username = username
    user.email = email
    user.phone_number = phone
    user.password = password
    user.save()
    return JsonResponse(response)


def deleteUserRecord(request):
    response = {"code": 2000, "message": "success"}
    record_id = request.GET.get("record_id")
    record = UserRecord.objects.get(id=record_id)
    datetime = record.time
    username = record.user.username
    medium_file_path = os.path.join(MEDIUM_PATH, username, datetime)
    delete_directory(medium_file_path)
    introduction_path = os.path.join(BASE_DIR, "Introduction", username, datetime)
    delete_directory(introduction_path)
    screenshot_path = os.path.join(BASE_DIR, "screenshot", username, datetime)
    delete_directory(screenshot_path)
    static_path = os.path.join(BASE_DIR, "static", username, datetime)
    delete_directory(static_path)
    record.delete()
    return JsonResponse(response)


@require_http_methods(["GET"])
def pdfDownload(request):
    user_id = request.GET.get("user_id")
    time = request.GET.get("time")
    username = UserProfile.objects.get(id=user_id).username
    static_file_path = os.path.join(BASE_DIR, "static", username, time, "ultimate_file.pdf")
    file = open(static_file_path, "rb")
    response = FileResponse(file)
    response["Content-Disposition"] = f"attachment; filename='{username}'+'{time}'+'.pdf"
    return response


@require_http_methods(["GET"])
def txtDownload(request):
    user_id = request.GET.get("user_id")
    time = request.GET.get("time")
    username = UserProfile.objects.get(id=user_id).username
    static_file_path = os.path.join(BASE_DIR, "static", username, time, "expanded_description.txt")
    file = open(static_file_path, "rb")
    response = FileResponse(file)
    response["Content-Disposition"] = f"attachment; filename='{username}'+'{time}'+'.txt"
    return response


@require_http_methods(["GET"])
def wordDownload(request):
    user_id = request.GET.get("user_id")
    time = request.GET.get("time")
    username = UserProfile.objects.get(id=user_id).username
    static_file_path = os.path.join(BASE_DIR, "static", username, time, "template_manual.docx")
    file = open(static_file_path, "rb")
    response = FileResponse(file)
    response["Content-Disposition"] = f"attachment; filename='{username}'+'{time}'+'.docx"
    return response


@require_http_methods(["GET"])
def getUserAllRecord(request):
    response = {"code": 0, "message": "success"}
    user_id = request.GET.get("user_id")
    records = UserRecord.objects.filter(user=UserProfile(id=user_id))
    response["data"] = []
    for r in records:
        dict_record = to_dict(r)
        record_id = dict_record["id"]
        record_item = UserRecord.objects.get(id=record_id)
        if dict_record.get("pdf_status", 0) == 0:
            relative_path = os.path.join(BASE_DIR, *dict_record["pdf_download"].split("/")[3:])
            if os.path.exists(relative_path):
                dict_record["pdf_status"] = 1
                record_item.pdf_status = 1
                record_item.save()
        if dict_record.get("code_status", 0) == 0:
            relative_path = os.path.join(BASE_DIR, *dict_record["code_download"].split("/")[3:])
            if os.path.exists(relative_path):
                dict_record["code_status"] = 1
                record_item.code_status = 1
                record_item.save()
        if dict_record.get("introduce_status", 0) == 0:
            relative_path = os.path.join(BASE_DIR, *dict_record["introduce_download"].split("/")[3:])
            if os.path.exists(relative_path):
                dict_record["introduce_status"] = 1
                record_item.introduce_status = 1
                record_item.save()
        response["data"].append(dict_record)
    return JsonResponse(response)


@require_http_methods(["POST"])
def register(request):
    response = {"code": 0, "message": "success"}
    data = json.loads(request.body)
    username = data.get("username")
    password = data.get("password")
    email = data.get("email")
    phone = data.get("phone")
    user_count = UserProfile.objects.filter(username=username).filter().count()
    if user_count != 0:
        response["code"] = 1
        response["message"] = "Username has already existed"
        return JsonResponse(response)
    else:
        user = UserProfile(username=username, password=password, email=email, phone_number=phone)
        user.save()
        return JsonResponse(response)


@require_http_methods(["GET"])
def getNameConfig(request):
    username = request.GET.get("username")
    datetime = request.GET.get("datetime")
    data = read_json(username, datetime)
    platform = data["platform"]
    response = {"code": 0, "message": "success"}
    response["name"] = platform
    return JsonResponse(response)


@require_http_methods(["GET"])
def getMenuConfig(request): # 获取侧边栏信息
    username = request.GET.get("username")
    datetime = request.GET.get("datetime")
    response = {"code": 0, "message": "success"}
    
    # 直接读取已生成的菜单配置
    json_file = os.path.join(MEDIUM_PATH, username, datetime, "menu.json")
    try:
        with open(json_file, "r", encoding="utf-8") as f:
            MENU_CONFIG = json.load(f)
        response["menu"] = MENU_CONFIG
    except Exception as e:
        print(f"读取菜单配置失败: {e}")
        response = {"code": 1, "message": "读取菜单配置失败"}
    
    return JsonResponse(response)



@require_http_methods(["GET"])
def getPageInfo(request): # 根据前端返回的当前点击的侧边栏id生成具体的页面代码
    username = request.GET.get("username")
    datetime = request.GET.get("datetime")
    colors = request.GET.get("colors")
    content_save = ""
    # 若 colors 为空，则使用默认值
    if not colors:
        color_list = ["#40c9c6", "#40c9c6", "#40c9c6", "#40c9c6"]
    else:
        color_list = colors.split(',')

    data = read_json(username, datetime)
    platform = data["platform"]
    response = {"code": 0, "message": "success"}

    # 读取侧边栏信息
    menu_file = os.path.join(MEDIUM_PATH, username, datetime, "menu.json")
    with open(menu_file, "r", encoding="utf-8") as f:
        menu_info = json.load(f)

    menu_item = request.GET.get("index")
    parent_idx, child_idx = menu_item.split("-")
    parent_idx = int(parent_idx)
    child_idx = int(child_idx)
    parent_name = list(menu_info.keys())[parent_idx - 1]
    child_name = menu_info[parent_name][child_idx - 1]

    # 读取组件文件内容
    component_files = [
        "SearchForm.vue",
        "BaseTable.vue",
        "EChartLine.vue",
        "EChartBar.vue",
        "EChartPie.vue",
        "Card.vue",
        "Leaderboard.vue",
    ]
    base_dir = os.path.join(os.path.dirname(__file__), "templates_components")
    components_code_list = []
    for filename in component_files:
        file_path = os.path.join(base_dir, filename)
        if os.path.exists(file_path):
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            components_code_list.append(content)
        else:
            print(f"[WARNING] File not found: {file_path}")
    all_components_code = "\n".join(components_code_list)

    # 第一次调用：生成页面代码
    MESSAGE = [{"role": "system", "content": "You are a helpful programmer and product manager"}]
    question = f"""
        现在有一个后台管理系统叫 {platform}，请为侧边栏父目录名称为 {parent_name} 下的 {child_name} 子目录所对应的页面设计 Vue2 版本的 Vue.js 页面代码，满足如下要求：

        【基本要求，必须遵守】
        1. 回复必须只包含 Vue 代码,请以文本的形式返回，必须以 "<template>" 开头，以 "</script>" 结尾，不要包含任何解释说明。
        2. 项目基于 Vue2 + Element-UI，禁止出现 Vue3 语法（如 setup、v-slot、#slot、Composition API 等），仅使用 Vue2 标准写法。
        3. 图表部分必须使用全局挂载的 echarts，通过 const chart = this.$echarts.init() 方式直接调用，不允许使用 import 语法，也不允许注册组件。
        4. 所有组件内部均应包含完整虚拟数据，且每个数据列表需包含 6-8 条真实业务场景示例数据，不从父组件传递 props。
        5. 表单字段必须分别使用单独变量进行 v-model 绑定，不要使用统一 formData 对象。
        6. Vue2 要求 template 中只能有一个根节点，请将页面所有内容统一包裹在一个根容器（如 <div>）。
        7. 页面中的所有按钮统一使用 Element-UI 默认样式，不需设置颜色，新增按钮需设置 class="add"，并绑定新增表单。
        8. el-table 表格的最后一列必须设置固定宽度（建议 width: 240px），确保“编辑”、“删除”、“详情”三个按钮在一行内横向并排展示，不得出现换行或错位现象。

        【UI设计与样式要求】
        1. 页面整体主题色为 {color_list}，在非按钮样式中尽可能多体现该色系，例如表头颜色、标签、进度条、Tabs 等，但按钮组件不需要设置颜色，保留 Element-UI 原本 type 的颜色。
        2. 所有模块（如表格、图表、搜索栏、统计卡片、Tabs 等）必须包裹在 el-card 组件内，el-card 设置 shadow="always"，并设置合理 margin 和 padding。
        3. 页面布局采用 el-row + el-col 组合，统计卡片和图表均一行两个，el-col 设置 span=12，避免单列堆叠。
        4. 表格使用 el-table，表头背景色必须为 {color_list[0]}，通过 header-cell-style 设置。表格最后一列包含“编辑”、“删除”、“详情”三个操作按钮，按钮需在一行排列，删除按钮带确认提示框，整个操作列必须通过设置 width 固定宽度，避免按钮换行。
        5. 分页使用 el-pagination，右对齐，风格统一。
        6. 顶部搜索栏用 el-card 包裹，包含 el-input、el-select 和 class="add" 的新增按钮。

        【样式布局要求】
        1. 全局样式注入（需在 <style> 中添加）：
           .el-table th {{ background-color: {color_list[0]} !important; }}
        2. 表格所在区域必须用 div 容器包裹，并设置 style="overflow-x: auto"，避免表格列过多导致内容溢出不可见。
        3. 所有 el-table 设置 style="width: 100%;"，建议开启 table-layout="auto"，让列宽根据内容自动撑开，仅操作按钮列设置固定宽度。
        4. 表格的“操作”按钮列必须设置 width="240"，其他所有 el-table-column 禁止设置 width，避免表格不能自适应满铺。

        【图表要求】
        1. 图表必须通过 this.$echarts.init() 初始化，禁止 import。
        2. 配置项必须完整，series 内必须包含 type、data、radius 等，禁止空图。
        3. 图表必须包含真实业务数据，放置在 el-card 中展示，布局合理美观。
        4. 所有图表的颜色（如 series 中的 itemStyle.color、lineStyle.color、bar 的柱状图颜色等），必须从{color_list}中选取颜色值，不能使用默认配色或其他色系。

        【组件丰富性要求】
        页面需合理引入至少 4 种 Element-UI 组件，颜色和样式统一，组件包括但不限于：
        - el-tabs、el-tag、el-switch、el-progress、el-divider、el-avatar、el-tooltip、el-breadcrumb、el-collapse、el-alert 等。

        【差异化要求】
        1. 主页页面必须包含至少一个 ECharts 图表，且布局与其他页面不同。
        2. 非主页页面需根据 {child_name} 的业务内容，自行设计布局，避免模板化。
        3. 示例数据必须真实可信，用户管理应有姓名、角色、手机号等字段，订单管理应有编号、金额、时间等信息。

        【弹窗要求】
        可以打开新增弹窗表单的按钮设置class为add，其余按钮不需要设置class。
        表单不设置title属性
        1. 弹窗包括自定义头部，新增表单和操作按钮 
        2. 弹窗的自定义头部，内容为新增XXX，设置slot="title"，覆盖本身标题,背景色必须设置为{color_list[0]}，自定义头部文字居中，
        3. 新增数据弹窗的表单需包含至少三个字段，字段类型覆盖文本输入、数字输入、下拉选择等，字段最好和页面表格对应
        4. 弹窗内必须包含“取消”和“确定”两个按钮，取消按钮 class="cancel"，点击取消关闭弹窗，操作按钮要符合主题色{color_list}色系中较为柔和的颜色
        5. 弹窗需出现在屏幕正中央，align-center 为 true，弹窗宽度为屏幕宽度的 30%。
        
        【已知 Bug 避免规则】
        1. el-table 表格必须设置 style="width: 100%;"，外层包裹 div 设置 style="overflow-x: auto;"。
        2. 操作按钮列使用 el-table-column width="240"，确保“编辑/删除/详情”按钮一行展示，禁止换行。
        3. 所有 el-card 外层容器需设置 margin/padding，避免表格压边显示异常。
        4. 弹窗需使用 align-center，且 :width="'30%'"，防止弹窗偏移或过宽，按钮保持一行横向排列。
        5. 所有 el-table-column 禁止设置固定 width，除了“操作列”；确保表格能自适应撑满表格区域。


        【已有模板提示】
        已有以下组件模板：{all_components_code}，你可以直接复制模板中的代码来复用，但禁止使用 import 语法，也不要直接注册组件调用。


    """
    MESSAGE.append({"role": "user", "content": question})


    try:
        page_code = api_call(MESSAGE)
    except Exception as e:
        print(f"所有尝试均失败: {e}")
        return JsonResponse({"code": 1, "message": "模型调用失败"})
    response["content"] = page_code
    content_save+=page_code
    MESSAGE.append({"role": "assistant", "content": page_code})
        # 保存前端代码到Word文档
    try:
        # 创建webfront文件夹
        webfront_path = os.path.join(MEDIUM_PATH, username, datetime, "webfront")
        if not os.path.exists(webfront_path):
            os.makedirs(webfront_path)
            
        # 处理前端代码，只保留template部分并删除空行
        processed_code = process_frontend_code(content_save)
            
        doc = Document()
         # 添加文件名作为标题
        doc.add_heading(child_name+"页面", level=1)
        doc.add_paragraph(processed_code)
        frontend_doc_path = os.path.join(MEDIUM_PATH, username, datetime, "webfront", f"前端_code_{menu_item}.docx")
        doc.save(frontend_doc_path)
        print(f"前端代码已保存到: {frontend_doc_path}")
    except Exception as e:
        print(f"保存前端代码到Word文档时出错: {e}")

    # 第二次调用：生成说明文档
    question = f"""
   请根据你所生成的页面代码，详细介绍该页面的业务功能与使用说明，要求如下：

    1. 回复必须以 ***** 开头；
    2. 换行进行说明内容
    3. 说明部分仅使用文字描述，不得包含任何代码格式（包括 markdown）；
    4. 回复说明内容格式如下：
    *****(换行)
    {{ {child_name} }}（换行）
    {{ 说明文字 }}
    5. 请用一段流畅自然的语言说明页面用途、用户操作流程及数据交互，不要分条、不使用编号；
    6. 字数在100字左右；
    7. 仅描述该页面能实现哪些核心功能，用户在页面中的操作逻辑；
    8. 不要出现与主题色、颜色代码、样式细节、像素占比等相关的描述,不需要对布局进行描述；
    9. 避免使用“页面整体视觉风格”、“界面美观性”、“卡片边框阴影”等表述；
    10. 若页面代码中包含弹窗，请将弹窗相关功能作为说明的最后一段单独说明。
    11. 说明文字中必须使用中文标点符号，严禁使用英文标点。
    """
    MESSAGE.append({"role": "user", "content": question})
    try:
        description = api_call(MESSAGE)
    except Exception as e:
        print(f"所有尝试均失败: {e}")
        return JsonResponse({"code": 1, "message": "模型调用失败"})
    content_save+=description
    # 保存说明文档
    txt_path = os.path.join(BASE_DIR, "Introduction", username, datetime)
    if not os.path.exists(txt_path):
        os.makedirs(txt_path)
    with open(os.path.join(txt_path, menu_item + ".txt"), "w", encoding="utf-8") as f:
        f.write(content_save)
    print(f"说明文档 {menu_item} 已保存")
    return JsonResponse(response)


@require_http_methods(["GET"])   # 生成登录界面信息
def getPageMain(request):
    import codecs  # 兼容不同编码

    username = request.GET.get("username")
    datetime = request.GET.get("datetime")
    colors = request.GET.get("colors")
    safe_datetime = quote(datetime)

    #background_url = build_static_url(username, datetime, "background.jpg")

    # 若 colors 为空，则使用默认值
    if not colors:
        color_list = ["#40c9c6", "#40c9c6", "#40c9c6", "#40c9c6"]
    else:
        color_list = colors.split(',')
    data = read_json(username, datetime)
    platform = data["platform"]
    response = {"code": 0, "message": "success"}
    reasoning_content = ""  # 定义完整思考过程
    content = ""     # 修改这里，使用 content 而不是 answer_content
    is_answering = False   # 判断是否结束思考过程并开始回复
    messages = []

    template_path = os.path.join(BASE_DIR, "llmserver", "templates_components", "tempLogin.vue")
    try:
        with codecs.open(template_path, 'r', encoding='utf-8') as f:
            vue_template_code = f.read()
    except Exception as e:
        vue_template_code = ""
        print(f"读取模板失败: {e}")

    # 定义登录的两种 Prompt
    question1_option1 = f"""
            现在有一个后台管理系统叫{platform}, 请为其用户登录页面生成所对应的基于Vue.js的代码, 满足如下要求:
            回复只需要以"<template>"开头, 以"<\/script>"结尾, 不需要其他任何解释说明内容.
            1.要求符合日常用户登录界面业务逻辑.
            2.组件基于element-ui框架.
            3.根容器背景使用指定图片URL：{IMG_URL}/{username}/{safe_datetime}/background.jpg，请确保图片可以显示,正确保留url中的空格，使用行内式background-image: url('{IMG_URL}/{username}/{safe_datetime}/background.jpg')
            确保图片可以完整显示，正好铺满（解决图片过大或者过小显示不全的问题）。
            4.卡片容器需要始终居中,样式要求如下width: 20%;position: absolute;top: 50%;left: 50%;transform: translate(-50%, -50%);
            5.登录卡片内部划分为标题区、输入区、操作按钮区三个纵向模块，各模块间距比例为1:2:1。
            6.标题区域固定于卡片顶部，显示欢迎登录等字样，使用大字号粗体文本居中显示，下方保留与输入区的动态间距。
            7.输入区域包含三个组件：
            ▸ 用户名输入框：标签+输入框
            ▸ 密码输入框：标签+输入框
            ▸ 验证码区，
            三个区域之间有些许间隔。
            8.操作按钮区包含一个按钮，按钮文本为"登录"，按钮颜色为蓝色，按钮宽度占满容器宽度，按钮下靠右有注册账号的灰色字样。
            9.要求只返回页面代码部分, 代码只包含template部分和script部分, 不需要style部分. 最重要的是以"<template>"开头, 以"<\/script>"结尾, 不需要其他任何解释说明内容。
        """

    question1_option2 = f"""
        现在有一个后台管理系统叫 {platform}，请为其用户登录页面生成对应的基于 Vue.js 的代码，满足如下要求：
        1. 只返回页面代码部分，代码只包含 <template> 到 </script>，不需要任何解释说明，也不需要 style 块。
        2. 所有样式均使用行内式写在元素的 style 属性中，不能使用 <style> 或 scoped 样式。
        3. 整个页面作为背景，宽高 100%（width: 100%; height: 100vh;），使用 background-size: cover 显示图片：url('{IMG_URL}/{username}/{safe_datetime}/background.jpg')，居中对齐但在右侧留出一块区域放置登录卡片（margin-right: 80px）。
        4. 登录卡片宽度 360px，白色背景，圆角，box-shadow，内部使用 Element-UI 的 <el-form>、<el-form-item>、<el-input>、<el-button> 组件，包含用户名、密码两个输入框，均带 prefix-icon。
        5. 登录按钮使用动态行内样式实现鼠标悬停时的背景渐变方向切换、轻微上移等效果。点击后在 onSubmit 中以 alert() 方式演示获取到的用户名、密码。
        6. 登录卡片内的标题需显示“欢迎登录 {platform}”等字样，请在 <h2> 标签中动态拼接 {platform}。
        7. 代码需符合常规登录业务逻辑示例：onSubmit 中可以先用 alert() 演示用户名、密码获取。
        8. 最终输出必须从 <template> 开始，到 </script> 结束，中间不能有额外文字或说明。
        9. 参考以下已有 Vue 模板结构，你只能直接复制其中的代码使用，不允许使用 import 语句：
        {vue_template_code}

        请严格遵循以上要求，仅返回符合条件的 Vue2 代码。
        """

    # 2. 根据全局随机值决定最终要用哪个 Prompt
    if RANDOM_CHOICE == 1:
        question1 = question1_option1
    else:
        question1 = question1_option2

    question2 = f"""
       请根据你所生成的页面代码，撰写一段流畅自然的文字，详细介绍该页面的核心业务功能和用户使用流程，要求如下：

       1. 回复必须以 ***** 开头，之后换行开始说明内容；
       2. 说明部分仅使用文字描述，不得包含任何代码格式（包括 markdown、代码片段等）；
       3. 回复说明内容格式如下：
        *****(换行)
       用户登录(换行)
       {{ 说明文字 }}
       4. 说明必须以一整段连贯文字进行撰写，避免使用编号、条列、分点描述等风格；
       5. 内容字数请在200字内；
       6. 请重点描述该页面实现了哪些功能、用户在页面中如何完成具体任务，前后端之间如何协同完成登录流程；
       7. 请避免出现页面配色、字体、大小、渐变、边框阴影等 UI 外观描述；
       8. 请勿撰写任何开发建议、技术扩展、接口设计说明或代码结构的说明；
       9. 说明文字中必须使用中文标点符号，严禁使用英文标点。
       """
    prompt = f"""现在有一个网页系统叫{platform},请帮其生成一张可以在登陆页面使用的背景图片,图片颜色符合{color_list[0]}的色系，图片内容和{platform}相关，不包含文字"""


    #print('----sync call, please wait a moment----')
    rsp = ImageSynthesis.call(
        api_key=api_keynum,
        model="wanx2.0-t2i-turbo",  # 使用较轻量的模型
        prompt=prompt,
        n=1,
        size='1024*576'  # 降低分辨率到 512*512
    )
    #print('response: %s' % rsp)
    if rsp.status_code == HTTPStatus.OK:
        for result in rsp.output.results:
            # 创建目录
            dir_path = os.path.join(BASE_DIR, "static", username, datetime)
            os.makedirs(dir_path, exist_ok=True)

            # 添加文件名
            file_name = os.path.join(dir_path, "background.jpg")  # 添加具体的文件名

            print(f"保存图片到: {file_name}")
            with open(file_name, 'wb+') as f:
                f.write(requests.get(result.url).content)
    else:
        print('sync_call Failed, status_code: %s, code: %s, message: %s' %
              (rsp.status_code, rsp.code, rsp.message))

    user_msg = {"role": "user", "content": question1}
    messages.append(user_msg)
    completion = client.chat.completions.create(
        model="qwq-plus-latest",  # 此处以 qwq-32b 为例，可按需更换模型名称
        messages=messages,
        # QwQ 模型仅支持流式输出方式调用
        stream=True
    )
    #print("\n" + "=" * 20 + "思考过程" + "=" * 20 + "\n")
    for chunk in completion:
        if not chunk.choices:
            print("\nUsage:")
            #print(chunk.usage)
        else:
            delta = chunk.choices[0].delta
            if hasattr(delta, 'reasoning_content') and delta.reasoning_content != None:
                #print(delta.reasoning_content, end='', flush=True)
                reasoning_content += delta.reasoning_content
            else:
                if delta.content != "" and is_answering is False:
                    #print("\n" + "=" * 20 + "完整回复" + "=" * 20 + "\n")
                    is_answering = True
                #print(delta.content, end='', flush=True)
                content += delta.content
    response["content"] = content
    messages.append({"role": "assistant", "content": content})
    user_msg = {"role": "user", "content": question2}
    messages.append(user_msg)
    content+="\n"
    completion = client.chat.completions.create(
        model="qwq-plus-latest",  # 此处以 qwq-32b 为例，可按需更换模型名称
        messages=messages,
        # QwQ 模型仅支持流式输出方式调用
        stream=True
    )
    #print("\n" + "=" * 20 + "思考过程" + "=" * 20 + "\n")
    for chunk in completion:
        if not chunk.choices:
            print("\nUsage:")
            #print(chunk.usage)
        else:
            delta = chunk.choices[0].delta
            if hasattr(delta, 'reasoning_content') and delta.reasoning_content != None:
                #print(delta.reasoning_content, end='', flush=True)
                reasoning_content += delta.reasoning_content
            else:
                if delta.content != "" and is_answering is False:
                    #print("\n" + "=" * 20 + "完整回复" + "=" * 20 + "\n")
                    is_answering = True
                #print(delta.content, end='', flush=True)
                content += delta.content
    txt_path = os.path.join(BASE_DIR, "Introduction", username, datetime)
    if not os.path.exists(txt_path):
        os.makedirs(txt_path)

    with open(os.path.join(txt_path, "0-0.txt"), "w", encoding="utf-8") as f:
        f.write(content)
        f.close()
    return JsonResponse(response)

@require_http_methods(["GET"])  # 生成注册页面信息
def getPageVice(request):
    import codecs  # 兼容不同编码

    username = request.GET.get("username")
    datetime = request.GET.get("datetime")
    safe_datetime = quote(datetime)
    data = read_json(username, datetime)
    platform = data["platform"]
    response = {"code": 0, "message": "success"}

    #background_url = build_static_url(username, datetime, "background.jpg")

    reasoning_content = ""  # 定义完整思考过程
    content = ""     # 修改这里，使用 content 而不是 answer_content
    is_answering = False   # 判断是否结束思考过程并开始回复
    messages = []

    template_path = os.path.join(BASE_DIR, "llmserver", "templates_components", "tempRes.vue")
    try:
        with codecs.open(template_path, 'r', encoding='utf-8') as f:
            vue_template_code = f.read()
    except Exception as e:
        vue_template_code = ""
        print(f"读取模板失败: {e}")

    # 定义注册的两种 Prompt
    question1_option1 = f"""
        现在有一个后台管理系统叫{platform}, 请为其用户注册页面生成所对应的基于Vue.js的代码, 满足如下要求:
        回复只需要以"<template>"开头, 以"<\/script>"结尾, 不需要其他任何解释说明内容.
        1.要求符合日常用户注册界面业务逻辑.
        2.组件基于element-ui框架.
        3.根容器背景使用指定图片URL：{IMG_URL}/{username}/{safe_datetime}/background.jpg，请确保图片可以显示,正确保留url中的空格，注意图片需要完整显示background-size: cover，使用行内式background-image: url('{IMG_URL}/{username}/{safe_datetime}/background.jpg')
        4.卡片容器需要始终居中,样式要求如下width: 20%;position: absolute;top: 50%;left: 50%;transform: translate(-50%, -50%);
        5.登录卡片内部划分为标题区、输入区、操作按钮区三个纵向模块，各模块间距比例为1:2:1。
        6.标题区域固定于卡片顶部，显示欢迎登录等字样，使用大字号粗体文本居中显示，下方保留与输入区的动态间距。
        7.输入区域项目自定义（例如：用户名，密码，邮箱，电话号码，确认密码。尽量大于五项），但需要标签+输入框的格式，可在下面加小字提示输入格式
        8.要求只返回页面代码部分, 代码只包含template部分和script部分, 不需要style部分. 最重要的是以"<template>"开头, 以"<\/script>"结尾, 不需要其他任何解释说明内容.
    """

    question1_option2 = f"""
    现在有一个后台管理系统叫 {platform}，请为其用户注册页面生成对应的基于 Vue.js 的代码，满足如下要求：
    1. 只返回页面代码部分，代码只包含 <template> 到 </script>，不需要任何解释说明内容，也不需要 style 块。
    2. 所有样式均使用行内式写在元素的 style 属性中，不能使用 <style> 或 scoped 样式。
    3. 页面背景使用图片：url('{IMG_URL}/{username}/{safe_datetime}/background.jpg')，要求覆盖全屏 (width: 100%; height: 100vh; background-size: cover;)，且在右侧留出一块区域放置注册卡片 (margin-right: 80px)。
    4. 注册卡片宽度 360px、白色背景、圆角、阴影、padding: 40px，内部包含用户名、密码、确认密码、邮箱、电话 5 个输入项，使用 label+input 行内样式加上小字提示 (small 标签)。最后有一个 <el-button> type="primary" 作为“注册”按钮，点击触发 onRegister 方法，方法内通过 alert() 显示表单内容。
    5. 代码需符合常规注册业务逻辑示例：onRegister 中可以先用 alert() 演示获取到的表单信息。
    6. 最终输出必须从 <template> 开始，到 </script> 结束，中间不能有任何额外文字或说明。
    7. 参考以下已有 Vue 模板结构，你只能直接复制其中的代码使用，不允许使用 import 语句：
    {vue_template_code}

    请严格遵循以上要求，仅返回符合条件的 Vue2 代码。
    """

    question2 = f"""
        请根据你所生成的页面代码，撰写一段自然流畅的文字，介绍该页面的核心业务功能和用户使用流程，要求如下：

        1. 回复必须以 ***** 开头，之后换行开始说明内容；
        2. 说明部分仅使用文字描述，不得包含任何代码格式（包括 markdown、代码片段等）；
        3. 回复说明内容格式如下：
         *****(换行)
        用户注册(换行)
        {{ 说明文字 }}
        4. 整体说明需以一段完整、连贯的叙述形式呈现，禁止使用编号、条列或分点描述；
        5. 内容字数请在200字内；
        6. 请围绕注册页面的核心功能展开描述，包括用户填写哪些信息、表单如何提交、验证机制如何、交互流程如何进行等；
        7. 请避免描述与页面外观有关的内容（如颜色、渐变、字体、布局样式等）；
        8. 不要写开发建议、接口说明、组件调用方式等技术实现细节；
        9. 说明文字中必须使用中文标点符号，严禁使用英文标点。
        """

    # 3. 根据同一个全局随机值来决定使用哪一个注册 Prompt
    if RANDOM_CHOICE == 1:
        question1 = question1_option1
    else:
        question1 = question1_option2

    user_msg = {"role": "user", "content": question1}
    messages.append(user_msg)
    completion = client.chat.completions.create(
        model="qwq-plus-latest",  # 此处以 qwq-32b 为例，可按需更换模型名称
        messages=messages,
        # QwQ 模型仅支持流式输出方式调用
        stream=True
    )
    #print("\n" + "=" * 20 + "思考过程" + "=" * 20 + "\n")
    for chunk in completion:
        if not chunk.choices:
            print("\nUsage:")
            #print(chunk.usage)
        else:
            delta = chunk.choices[0].delta
            if hasattr(delta, 'reasoning_content') and delta.reasoning_content != None:
                #print(delta.reasoning_content, end='', flush=True)
                reasoning_content += delta.reasoning_content
            else:
                if delta.content != "" and is_answering is False:
                    #print("\n" + "=" * 20 + "完整回复" + "=" * 20 + "\n")
                    is_answering = True
                #print(delta.content, end='', flush=True)
                content += delta.content
    response["content"] = content
   
    messages.append({"role": "assistant", "content": content})
    user_msg = {"role": "user", "content": question2}
    messages.append(user_msg)
    content+="\n"
    completion = client.chat.completions.create(
        model="qwq-plus-latest",  # 此处以 qwq-32b 为例，可按需更换模型名称
        messages=messages,
        # QwQ 模型仅支持流式输出方式调用
        stream=True
    )
    #print("\n" + "=" * 20 + "思考过程" + "=" * 20 + "\n")
    for chunk in completion:
        if not chunk.choices:
            print("\nUsage:")
            #print(chunk.usage)
        else:
            delta = chunk.choices[0].delta
            if hasattr(delta, 'reasoning_content') and delta.reasoning_content != None:
                #print(delta.reasoning_content, end='', flush=True)
                reasoning_content += delta.reasoning_content
            else:
                if delta.content != "" and is_answering is False:
                    #print("\n" + "=" * 20 + "完整回复" + "=" * 20 + "\n")
                    is_answering = True
                #print(delta.content, end='', flush=True)
                content += delta.content
    txt_path = os.path.join(BASE_DIR, "Introduction", username, datetime)
    if not os.path.exists(txt_path):
        os.makedirs(txt_path)

    with open(os.path.join(txt_path, "0-1.txt"), "w", encoding="utf-8") as f:
        f.write(content)
        f.close()
    return JsonResponse(response)



@require_http_methods(["GET"])
def registerDownloadDocx(request):
    user_id = request.GET.get("user_id")
    time = request.GET.get("time")
    username = UserProfile.objects.get(id=user_id).username
    static_file_path = os.path.join(BASE_DIR, "static", username, time, "软著注册表.txt")
    file = open(static_file_path, "rb")
    response = FileResponse(file)
    response["Content-Disposition"] = f"attachment; filename='{username}'+'{time}'+'.txt"
    return response

def run_app_script(software_name, programming_language, user_id, time_str):
    """运行app.py脚本生成软著注册表"""
    try:
        username = UserProfile.objects.get(id=user_id).username
        
        # 确保输出目录存在
        output_dir = os.path.join(BASE_DIR, "static", username, time_str)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        # 生成输出文件路径 - 使用txt作为主要格式
        output_path = os.path.join(output_dir, "软著注册表.txt")
        
        # 运行app.py脚本
        subprocess.run(
            ["python3", os.path.join(BASE_DIR, "llmserver", "components", "app.py"), 
             software_name, programming_language, output_path, output_dir],
            check=True
        )
        
        print(f"软著注册表生成成功: {output_path}")
    except Exception as e:
        print(f"运行app.py脚本出错: {e}")

def process_frontend_code(code):
    # 提取<template>和</template>之间的内容
    template_pattern = re.compile(r'<template>(.*?)</template>', re.DOTALL)
    template_match = template_pattern.search(code)
    
    if template_match:
        template_content = template_match.group(1)
        # 删除空行
        lines = template_content.split('\n')
        non_empty_lines = [line for line in lines if line.strip()]
        processed_content = '\n'.join(non_empty_lines)
        return processed_content
    
    # 如果没有找到template标签，则返回原始代码
    return code

@require_http_methods(["GET"])
def generateRegistration(request):
    """生成软著注册表"""
    response = {"code": 0, "message": "success"}
    try:
        user_id = request.GET.get("user_id")
        record_id = request.GET.get("record_id")
        
        # 从数据库中查询项目信息
        record = UserRecord.objects.get(id=record_id)
        software_name = record.name  # 获取软件名称
        programming_language = record.language  # 获取编程语言
        
        # 启动线程运行app.py
        thread = threading.Thread(
            target=run_app_script,
            args=(software_name, programming_language, user_id, record.time)
        )
        thread.daemon = True
        thread.start()
        
        # 准备路径
        username = UserProfile.objects.get(id=user_id).username
        
        # 设置软著注册表下载路径和状态
        record.register_download = f"{WEB_URL}/{username}/{record.time}/软著注册表.txt"
        record.register_status = 0
        record.save()
        
        # 返回成功消息
        response["message"] = "软著注册表正在生成，请稍后查看"
        return JsonResponse(response)
    except UserRecord.DoesNotExist:
        return JsonResponse({"code": 1, "message": "记录不存在"})
    except Exception as e:
        print(f"设置软著注册表信息出错: {e}")
        return JsonResponse({"code": 1, "message": f"设置软著注册表信息失败: {str(e)}"})
