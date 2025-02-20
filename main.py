import subprocess
import os
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import json
import ast

# Vue 项目路径
vue_project_path = r"E:\ui_generator"

TXT_PATH = 'backend/Introduction/'
IMAGE_PATH = 'screenshot'
if not os.path.exists(IMAGE_PATH):
    os.makedirs(IMAGE_PATH)

PLATFORM = '共享充电宝后台管理系统'

LLM_BASE_URL = 'https://dashscope.aliyuncs.com/compatible-mode/v1'
LLM_API_KEY = 'sk-8fafb7bba6d04e45a989191a8820f7a1'

client = OpenAI(
    api_key=LLM_API_KEY,
    base_url=LLM_BASE_URL
)

# 启动 Vue 项目
def start_vue_project():
    # 进入 Vue 项目目录并启动开发服务器
    process = subprocess.Popen(
        ["npm", "run", "serve"],
        cwd=vue_project_path,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        shell=True
    )
    print("Vue 项目已启动...")
    return process

# 关闭 Vue 项目
def stop_vue_project(process):
    process.terminate()
    print("Vue 项目已关闭。")

def generate_config():
    with open('backend/menu.json', 'r', encoding='utf-8') as f:
        menu = json.load(f)
        f.close()
    json_str = json.dumps(menu, ensure_ascii=False)
    MESSAGE = [{'role': 'system', 'content': 'You are a helpful programmer and product manager'}]
    question = f"""
    现在有一个名为{PLATFORM}的后台管理系统,其侧边栏为{json_str},请为其设计如下的一些配置信息:
    1.开发硬件环境,包括处理器,内存,硬盘,显卡等配置信息.
    2.运行硬件环境,包括处理器,内存,硬盘,带宽等配置信息.
    3.开发目的,说明文字不少于300字.
    4.功能描述,说明文字不少于600字.
    5.以Object形式返回，返回内容开始以'{{'开头，不需要额外任何解释说明.
    6.语言为中文
    7.配置信息为纯文字说明,不需要换行
    """
    query = {'role': 'user', 'content': question}
    MESSAGE.append(query)
    completion = client.chat.completions.create(
        model="qwen-coder-plus-latest",
        messages=MESSAGE
    )
    data_dict = json.loads(completion.model_dump_json())
    content = data_dict['choices'][0]['message']['content']
    return content


def take_screenshot(index, driver):
    script = f"""
        window.dataLoaded = false;
        window.addEventListener('data-loaded_{index}', () => {{
            window.dataLoaded = true;
        }});
        """
    driver.execute_script(script)
    wait = WebDriverWait(driver, 500)
    wait.until(lambda d: d.execute_script("return window.dataLoaded;"))
    time.sleep(2)

    screenshot_path = os.path.join(IMAGE_PATH, index+'.png')
    driver.save_screenshot(screenshot_path)
    print(f"截图已保存为 {screenshot_path}")


def main_process():
    chrome_driver_path = r"chromedriver-win64\chromedriver.exe"
    # 配置 Selenium 使用 Chrome 浏览器
    options = webdriver.ChromeOptions()
    options.add_experimental_option("detach", True)  # 让 Chrome 保持打开，不会自动关闭
    options.add_argument("--start-maximized")  # 启动时窗口最大化
    options.add_argument("--disable-blink-features=AutomationControlled")  # 规避检测
    # options.add_argument("--headless")  # 无头模式
    # options.add_argument("--disable-gpu")  # 禁用 GPU 加速
    # options.add_argument("--window-size=1280,800")  # 设置窗口大小

    # 启动浏览器
    driver = webdriver.Chrome(service=Service(chrome_driver_path), options=options)

    driver.get('http://localhost:8080/login')
    take_screenshot('0-0', driver)

    driver.get('http://localhost:8080/register')
    take_screenshot('0-1', driver)

    # 访问 Vue 项目
    driver.get("http://localhost:8080/")

    take_screenshot('1-1', driver)

    try:
        sidebarItems = driver.find_element(By.ID, 'menu')
        parent_items = sidebarItems.find_elements(By.CLASS_NAME, 'parent-menu')
        for index, parent in enumerate(parent_items):
            try:
                sub_menu = parent.find_element(By.XPATH, ".//following-sibling::ul")  # 这里假设子菜单是 <ul> 标签
                if not sub_menu.is_displayed():
                    parent.click()  # **如果子目录未显示，则点击父目录展开**
                    time.sleep(1)  # 等待展开
            except:
                print("无法找到子菜单，可能结构不同")

            sub_menus = parent.find_elements(By.XPATH, ".//following-sibling::ul//li")
            for sub_index, sub_item in enumerate(sub_menus):
                sub_item.click()
                menu_index = str(index+1)+'-'+str(sub_index+1)
                if menu_index != '1-1':
                    take_screenshot(menu_index, driver)


    except Exception as e:
        print("无法找到侧边栏目录:", e)

    # 关闭浏览器
    driver.quit()

# 主函数
def main():
    # 启动 Vue 项目
    # vue_process = start_vue_project()

    # 等待 Vue 项目启动
    # time.sleep(10)

    # try:
        # 截图
    main_process()
    # finally:
    #     # 关闭 Vue 项目
    #     stop_vue_project(vue_process)


def read_txt_file(filename):
    context = {}
    line_data = []
    flag = False
    index = 0
    with open(filename, 'r', encoding='utf-8') as fr:
        lines = fr.readlines()
        for line in lines:
            if line.__contains__('*****'):
                flag = True
                continue
            if flag:
                if index == 0:
                    context['name'] = line.replace('\n', '')+'页面'
                    index += 1
                else:
                    line_data.append(line)

    content = ''.join(line_data).replace('\n\n', '\n')
    context['content'] = content
    return context


def get_image_info(file):
    index = file.split('.')[0]
    for file in os.listdir(IMAGE_PATH):
        if file.split('.')[0] == index:
            return os.path.join(IMAGE_PATH, file)


def generate_word_template():
    doc = Document('template.docx')

    main_info = {
        'title': '共享充电宝后台管理系统V1.0',
        'table': [
            {'version': 'v1.0', 'date': '2025-02-18', 'name': '张三', 'info': '初始版本'},
            {'version': 'v1.1', 'date': '2025-03-15', 'name': '张三', 'info': '增加数据统计页面'},
        ]
    }

    config = [
        {
            'name': '环境描述',
            'subsection': []
        },
        {
            'name': '开发目的',
            'content': ''
        },
        {
            'name': '功能描述',
            'content': ''
        }
    ]

    main_content = {
        'title': '使用说明',
        'subsection': []
    }

    config_content = generate_config()
    config_content = ast.literal_eval(config_content)

    for i, key in enumerate(config_content.keys()):
        if i < 2:
            config[0]['subsection'].append({
                'name': key,
                'content': config_content[key]
            })
        else:
            config[i-1]['content'] = config_content[key]

    for paragraph in doc.paragraphs:
        if '{{ title }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ title }}', main_info['title'])

    table = doc.tables[0]
    for i, row_data in enumerate(main_info['table']):
        table.cell(i + 1, 0).text = row_data['version']
        table.cell(i + 1, 1).text = str(row_data['date'])
        table.cell(i + 1, 2).text = row_data['name']
        table.cell(i + 1, 3).text = row_data['info']

    for section in config:
        doc.add_page_break()
        doc.add_heading(section['name'], level=1)
        if 'subsection' in section.keys():
            for subsection in section['subsection']:
                doc.add_heading(subsection['name'], level=2)
                doc.add_paragraph(subsection['content'])
        else:
            doc.add_paragraph(section['content'])

    for file in os.listdir(TXT_PATH):
        filename = os.path.join(TXT_PATH, file)
        context = read_txt_file(filename)
        image = get_image_info(file)
        context['image'] = image
        main_content['subsection'].append(context)

    doc.add_page_break()
    doc.add_heading(main_content['title'], level=1)
    for section in main_content['subsection']:
        doc.add_heading(section['name'], level=2)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.add_run().add_picture(section['image'], width=Inches(5.5))
        doc.add_paragraph(section['content'])

    doc.save('template_manual.docx')


if __name__ == "__main__":
    # main()
    generate_word_template()